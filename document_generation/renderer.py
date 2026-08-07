from __future__ import annotations

from datetime import timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt

from .docx_helpers import (
    DEFAULT_SIGNATURE_LINE_SPACE_BEFORE_PT,
    add_signature_cell,
    apply_run_font,
    clear_cell,
    configure_layout_table,
    set_cell_border,
    set_table_borders,
    set_table_indent,
)
from .models import DocumentContext, Person
from .text_utils import (
    display_person_name,
    format_currency,
    format_date,
    honorific_name,
    normalize_address,
)


BLOCK_MARKERS = {
    "{{ block:director_fee_section }}",
    "{{ block:director_signatures }}",
    "{{ block:member_signatures }}",
    "{{ block:shareholder_addresses }}",
    "{{ block:acknowledgement_signatures }}",
    "{{ block:attendance_signatures }}",
}
SIGNATURE_GRID_TOTAL_REDUCTION = Mm(2)
FEE_CONTENT_INDENT = Cm(1)
FEE_TABLE_CELL_START_MARGIN_TWIPS = 60
FEE_NAME_COLUMN_INDENT = Pt(3)
ATTENDANCE_MINIMUM_BODY_ROWS = 10
ATTENDANCE_ROW_MINIMUM_HEIGHT = Inches(0.48)


def _iter_paragraphs(container):
    seen: set[object] = set()
    for paragraph in container.paragraphs:
        marker = paragraph._p
        if marker not in seen:
            seen.add(marker)
            yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in _iter_paragraphs(cell):
                    marker = paragraph._p
                    if marker not in seen:
                        seen.add(marker)
                        yield paragraph


def _replace_token(paragraph, token: str, replacement: str) -> None:
    while token in "".join(run.text for run in paragraph.runs):
        full_text = "".join(run.text for run in paragraph.runs)
        start = full_text.index(token)
        end = start + len(token)
        cursor = 0
        first_index = None
        last_index = None
        first_offset = 0
        last_offset = 0
        for index, run in enumerate(paragraph.runs):
            run_end = cursor + len(run.text)
            if first_index is None and start < run_end:
                first_index = index
                first_offset = start - cursor
            if end <= run_end:
                last_index = index
                last_offset = end - cursor
                break
            cursor = run_end
        if first_index is None or last_index is None:
            return
        first = paragraph.runs[first_index]
        last = paragraph.runs[last_index]
        if first_index == last_index:
            first.text = first.text[:first_offset] + replacement + first.text[last_offset:]
        else:
            first.text = first.text[:first_offset] + replacement
            for index in range(first_index + 1, last_index):
                paragraph.runs[index].text = ""
            last.text = last.text[last_offset:]


def _replace_scalars(document: Document, replacements: dict[str, str]) -> None:
    containers = [document]
    for section in document.sections:
        containers.extend((section.header, section.footer))
    for container in containers:
        for paragraph in _iter_paragraphs(container):
            for token, value in replacements.items():
                if token in paragraph.text:
                    _replace_token(paragraph, token, value)


def _find_marker(document: Document, marker: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == marker:
            return paragraph
    raise ValueError(f"Template block marker was not found: {marker}")


def _has_marker(document: Document, marker: str) -> bool:
    return any(paragraph.text.strip() == marker for paragraph in document.paragraphs)


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    parent.remove(paragraph._p)


def _paragraph_before(document: Document, marker, text: str = ""):
    paragraph = document.add_paragraph()
    marker._p.addprevious(paragraph._p)
    if text:
        apply_run_font(paragraph.add_run(text), 12)
    return paragraph


def _table_before(document: Document, marker, rows: int, cols: int):
    table = document.add_table(rows=rows, cols=cols)
    marker._p.addprevious(table._tbl)
    return table


def _insert_signature_grid(
    document: Document,
    marker_text: str,
    people: list[Person],
    prefix: str = "",
) -> None:
    marker = _find_marker(document, marker_text)
    row_count = (len(people) + 1) // 2
    line_space_before_pt = (
        DEFAULT_SIGNATURE_LINE_SPACE_BEFORE_PT
        - SIGNATURE_GRID_TOTAL_REDUCTION.pt / row_count
    )
    table = _table_before(document, marker, row_count, 3)
    configure_layout_table(table, (3.0, 0.65, 3.0))
    for row_index in range(row_count):
        left_index = row_index * 2
        right_index = left_index + 1
        add_signature_cell(
            table.cell(row_index, 0),
            people[left_index].name.upper(),
            prefix,
            line_space_before_pt,
        )
        clear_cell(table.cell(row_index, 1))
        if right_index < len(people):
            add_signature_cell(
                table.cell(row_index, 2),
                people[right_index].name.upper(),
                prefix,
                line_space_before_pt,
            )
        else:
            clear_cell(table.cell(row_index, 2))
    _remove_paragraph(marker)


def _insert_address_grid(document: Document, context: DocumentContext) -> None:
    marker = _find_marker(document, "{{ block:shareholder_addresses }}")
    row_count = (len(context.members) + 1) // 2
    table = _table_before(document, marker, row_count, 3)
    configure_layout_table(table, (3.0, 0.65, 3.0))
    for row_index in range(row_count):
        for column, person_index in ((0, row_index * 2), (2, row_index * 2 + 1)):
            cell = table.cell(row_index, column)
            clear_cell(cell)
            if person_index >= len(context.members):
                continue
            member = context.members[person_index]
            name_paragraph = cell.paragraphs[0]
            name_paragraph.paragraph_format.space_after = Pt(1)
            apply_run_font(name_paragraph.add_run(display_person_name(member.name)), 12)
            for line in normalize_address(member.address).splitlines():
                paragraph = cell.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                apply_run_font(paragraph.add_run(line), 12)
        clear_cell(table.cell(row_index, 1))
    _remove_paragraph(marker)


def _insert_attendance_table(document: Document, context: DocumentContext) -> None:
    """Insert the Section 90 NAME/SIGNATURE attendance register."""
    marker = _find_marker(document, "{{ block:attendance_signatures }}")
    body_row_count = max(ATTENDANCE_MINIMUM_BODY_ROWS, len(context.members))
    table = _table_before(document, marker, body_row_count + 1, 2)
    configure_layout_table(table, (3.45, 2.85))
    set_table_borders(table, "single")

    header_row = table.rows[0]
    header_row.height = ATTENDANCE_ROW_MINIMUM_HEIGHT
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    header_properties = header_row._tr.get_or_add_trPr()
    if header_properties.find(qn("w:tblHeader")) is None:
        header_properties.append(OxmlElement("w:tblHeader"))
    for column, label in enumerate(("NAME", "SIGNATURE")):
        cell = header_row.cells[column]
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if column == 0
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        apply_run_font(paragraph.add_run(label), 12, True)

    for row_index in range(1, body_row_count + 1):
        row = table.rows[row_index]
        row.height = ATTENDANCE_ROW_MINIMUM_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        name_cell, signature_cell = row.cells
        clear_cell(name_cell)
        clear_cell(signature_cell)
        name_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        signature_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if row_index <= len(context.members):
            apply_run_font(
                name_cell.paragraphs[0].add_run(
                    context.members[row_index - 1].name.upper()
                ),
                12,
            )
    _remove_paragraph(marker)


def _insert_fee_section(document: Document, context: DocumentContext) -> None:
    marker = _find_marker(document, "{{ block:director_fee_section }}")
    if not context.has_director_fee:
        _remove_paragraph(marker)
        return

    heading = _paragraph_before(document, marker)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.left_indent = FEE_CONTENT_INDENT
    heading.paragraph_format.first_line_indent = -FEE_CONTENT_INDENT
    heading.paragraph_format.tab_stops.add_tab_stop(FEE_CONTENT_INDENT)
    apply_run_font(heading.add_run("5.\t"), 12)
    apply_run_font(heading.add_run("PAYMENT OF DIRECTORS’ FEES"), 12, True)
    body = _paragraph_before(
        document,
        marker,
        (
            f"That the payment of {format_currency(context.director_fee_total)} "
            f"as Directors’ fees for the financial year ended "
            f"{format_date(context.financial_year_end)} be recommended to the "
            "members for their approval as follows:-"
        ),
    )
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.left_indent = FEE_CONTENT_INDENT
    body.paragraph_format.space_after = Pt(3)

    table = _table_before(
        document,
        marker,
        len(context.fee_allocations) + 2,
        2,
    )
    configure_layout_table(table, (3.3, 1.6))
    set_table_indent(
        table,
        FEE_CONTENT_INDENT.twips - FEE_TABLE_CELL_START_MARGIN_TWIPS,
    )
    for column, label in enumerate(("Name", "Amount")):
        cell = table.cell(0, column)
        clear_cell(cell)
        if column == 0:
            cell.paragraphs[0].paragraph_format.left_indent = FEE_NAME_COLUMN_INDENT
        apply_run_font(cell.paragraphs[0].add_run(label), 12, True)
    for row_index, allocation in enumerate(context.fee_allocations, start=1):
        name_cell = table.cell(row_index, 0)
        amount_cell = table.cell(row_index, 1)
        clear_cell(name_cell)
        clear_cell(amount_cell)
        name_cell.paragraphs[0].paragraph_format.left_indent = FEE_NAME_COLUMN_INDENT
        apply_run_font(
            name_cell.paragraphs[0].add_run(display_person_name(allocation.name)),
            12,
        )
        apply_run_font(
            amount_cell.paragraphs[0].add_run(format_currency(allocation.amount)),
            12,
        )
    total_row = len(context.fee_allocations) + 1
    clear_cell(table.cell(total_row, 0))
    total_cell = table.cell(total_row, 1)
    clear_cell(total_cell)
    apply_run_font(
        total_cell.paragraphs[0].add_run(format_currency(context.director_fee_total)),
        12,
    )
    set_cell_border(total_cell, "top", "single", "6")
    set_cell_border(total_cell, "bottom", "double", "6")
    _remove_paragraph(marker)


def _scalar_values(context: DocumentContext) -> dict[str, str]:
    signer_names = [
        honorific_name(person.name, person.gender)
        for person in context.statement_signers
    ]
    if len(signer_names) == 1:
        signers = signer_names[0]
    else:
        signers = " and ".join(signer_names)
    member_count = len(context.members)
    values = {
        "{{ company_name }}": context.company_name,
        "{{ registration_no }}": context.registration_no,
        "{{ dwr_clause }}": context.dwr_clause,
        "{{ mwr_clause }}": context.mwr_clause,
        "{{ financial_year_start }}": format_date(context.financial_year_start),
        "{{ financial_year_end }}": format_date(context.financial_year_end),
        "{{ financial_year_end_upper }}": format_date(context.financial_year_end).upper(),
        "{{ board_approval_date }}": format_date(context.board_approval_date),
        "{{ circulation_date }}": format_date(context.circulation_date),
        "{{ lapse_date }}": format_date(context.circulation_date + timedelta(days=28)),
        "{{ statement_signers_with_titles }}": signers,
        "{{ statement_signer_authority }}": (
            "be and is" if len(context.statement_signers) == 1 else "be and are"
        ),
        "{{ declarant_with_title }}": honorific_name(
            context.statutory_declarant.name,
            context.statutory_declarant.gender,
        ),
        "{{ declarant_name_upper }}": context.statutory_declarant.name.upper(),
        "{{ auditor_name }}": context.auditor_name,
        "{{ circulation_section_number }}": "6" if context.has_director_fee else "5",
        "{{ filing_section_number }}": "7" if context.has_director_fee else "6",
        "{{ director_signature_heading }}": (
            "SOLE DIRECTOR" if len(context.directors) == 1 else "BOARD OF DIRECTORS"
        ),
        "{{ member_signature_heading }}": "SOLE MEMBER" if member_count == 1 else "MEMBERS",
        "{{ shareholder_heading }}": (
            "Shareholder of Company" if member_count == 1 else "Shareholders of Company"
        ),
        "{{ salutation }}": "Dear Sir," if member_count == 1 else "Dear Sirs,",
    }
    if context.section90 is not None:
        details = context.section90
        retiring_names = [
            honorific_name(person.name, person.gender)
            for person in details.retiring_directors
        ]
        if len(retiring_names) == 1:
            retiring_text = retiring_names[0]
            retirement_pronoun = (
                "herself"
                if details.retiring_directors[0].gender.upper() == "FEMALE"
                else "himself"
            )
        else:
            retiring_text = ", ".join(retiring_names[:-1]) + " and " + retiring_names[-1]
            retirement_pronoun = "themselves"
        ordinal_prefix = f"{details.agm_ordinal} " if details.agm_ordinal else ""
        values.update(
            {
                "{{ agm_title }}": f"{ordinal_prefix}ANNUAL GENERAL MEETING",
                "{{ agm_title_sentence }}": f"{ordinal_prefix}Annual General Meeting",
                "{{ meeting_date }}": format_date(details.meeting_date),
                "{{ notice_date }}": format_date(details.notice_date),
                "{{ letter_date }}": format_date(details.letter_date),
                "{{ meeting_start_time }}": details.meeting_start_time,
                "{{ meeting_end_time }}": details.meeting_end_time,
                "{{ meeting_venue }}": "\n".join(details.venue_lines),
                "{{ meeting_venue_inline }}": " ".join(details.venue_lines),
                "{{ retiring_directors_with_titles }}": retiring_text,
                "{{ retiring_director_noun }}": (
                    "director" if len(retiring_names) == 1 else "directors"
                ),
                "{{ retiring_director_noun_upper }}": (
                    "DIRECTOR" if len(retiring_names) == 1 else "DIRECTORS"
                ),
                "{{ retiring_pronoun }}": retirement_pronoun,
                "{{ retiring_offer_verb }}": (
                    "offers" if len(retiring_names) == 1 else "offer"
                ),
                "{{ retiring_is_are }}": "is" if len(retiring_names) == 1 else "are",
                "{{ chair_with_title }}": honorific_name(
                    context.statutory_declarant.name,
                    context.statutory_declarant.gender,
                ),
                "{{ chair_name_upper }}": context.statutory_declarant.name.upper(),
            }
        )
    return values


def render_document(
    template_path: Path,
    output_path: Path,
    context: DocumentContext,
) -> Path:
    document = Document(str(template_path))
    if _has_marker(document, "{{ block:director_fee_section }}"):
        _insert_fee_section(document, context)
    if _has_marker(document, "{{ block:director_signatures }}"):
        _insert_signature_grid(
            document,
            "{{ block:director_signatures }}",
            context.directors,
        )
    if _has_marker(document, "{{ block:member_signatures }}"):
        _insert_signature_grid(
            document,
            "{{ block:member_signatures }}",
            context.members,
        )
    if _has_marker(document, "{{ block:attendance_signatures }}"):
        _insert_attendance_table(document, context)
    if _has_marker(document, "{{ block:shareholder_addresses }}"):
        _insert_address_grid(document, context)
    if _has_marker(document, "{{ block:acknowledgement_signatures }}"):
        _insert_signature_grid(
            document,
            "{{ block:acknowledgement_signatures }}",
            context.members,
            prefix="Name: ",
        )
    _replace_scalars(document, _scalar_values(context))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
