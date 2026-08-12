from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


HERE = Path(__file__).resolve().parent
TEMPLATES = HERE.parent / "templates" / "agm"
LEGACY_FIRST_AGM = TEMPLATES / "agm_approve_accounts_template 1.docx"
FIRST_AGM = TEMPLATES / "first_agm_approve_accounts_template.docx"
SECTION_90 = TEMPLATES / "agm_approve_accounts_template_section_90.docx"


def _set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _format_section(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.49)
    section.bottom_margin = Inches(0.28)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)


def _style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for section in document.sections:
        _format_section(section)


def _p(
    document: Document,
    text: str = "",
    *,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 2,
    indent: float = 0,
    first: float | None = None,
    size: float = 11,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.left_indent = Inches(indent)
    if first is not None:
        paragraph.paragraph_format.first_line_indent = Inches(first)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return paragraph


def _company_header(document: Document) -> None:
    _p(document, "{{ company_name }}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, size=12)
    _p(document, "Registration No. {{ registration_no }}", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    _p(document, "(Incorporated in Malaysia)", align=WD_ALIGN_PARAGRAPH.CENTER, after=7)


def _title(document: Document, text: str) -> None:
    paragraph = _p(document, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=7, size=12)
    paragraph.paragraph_format.keep_with_next = True


def _resolution(document: Document, number: int, heading: str, body: str) -> None:
    paragraph = _p(document, f"{number}.\t{heading}", bold=True, before=3, after=1)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.35))
    body_paragraph = _p(
        document,
        body,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=3,
        indent=0.35,
        first=-0.01,
    )
    body_paragraph.paragraph_format.keep_with_next = False


def _new_page(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _format_section(section)


def _build_section_90() -> None:
    document = Document()
    _style_document(document)

    _company_header(document)
    _title(document, "DIRECTORS’ WRITTEN RESOLUTION")
    _p(document, "{{ dwr_clause }}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    _p(document, "It is hereby resolved:-", after=4)
    _resolution(
        document,
        1,
        "DIRECTORS’ REPORT",
        "That the Directors’ Report, a copy of which is attached hereto, be approved and adopted as the Directors’ Report required pursuant to Section 252(2) of the Companies Act, 2016 to be made and attached to the Company’s Audited Financial Statements for the financial year ended {{ financial_year_end }}.",
    )
    _resolution(
        document,
        2,
        "AUDITED FINANCIAL STATEMENTS",
        "That the Statement of Financial Position as at {{ financial_year_end }} and Statement of Comprehensive Income for the financial year ended {{ financial_year_end }}, which have been made out in accordance with the applicable approved accounting standards, be and are hereby approved.",
    )
    _resolution(
        document,
        3,
        "SIGNING OF DIRECTORS’ REPORT AND STATEMENT",
        "That {{ statement_signers_with_titles }} {{ statement_signer_authority }} hereby authorised to sign for and on behalf of the Board, pursuant to Sections 252(2)(b), 251(2) and 251(3) of the Companies Act, 2016, the Directors’ Report referred to above and the Directors’ Statement accompanying the Audited Financial Statements for the financial year ended {{ financial_year_end }}.",
    )
    _resolution(
        document,
        4,
        "SIGNING OF STATUTORY DECLARATION",
        "That {{ declarant_with_title }} be and is hereby authorised to sign, pursuant to Section 251(1)(b) of the Companies Act, 2016, the Statutory Declaration accompanying the Audited Financial Statements for the financial year ended {{ financial_year_end }}.",
    )
    _resolution(
        document,
        5,
        "CONVENING OF {{ agm_title }}",
        "That the {{ agm_title_sentence }} of the Company be held and convened at {{ meeting_venue_inline }} on {{ meeting_date }} at {{ meeting_start_time }} and that any one of the Directors of the Company be authorised to issue the notice of meeting accordingly.",
    )
    _p(document, "{{ director_signature_heading }}", bold=True, before=5, after=0)
    _p(document, "{{ block:director_signatures }}", after=0)
    _p(document, "Date: {{ board_approval_date }}", before=2)

    _new_page(document)
    _company_header(document)
    _title(document, "NOTICE OF {{ agm_title }}")
    _p(
        document,
        "NOTICE IS HEREBY GIVEN THAT THE {{ agm_title }} OF THE COMPANY WILL BE HELD AT {{ meeting_venue_inline }} ON {{ meeting_date }} AT {{ meeting_start_time }} FOR THE FOLLOWING PURPOSES:-",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=12,
    )
    _title(document, "AGENDA")
    _p(document, "1.\tTo receive the Audited Financial Statements for the financial year ended {{ financial_year_end }} together with the Reports of Directors and Auditors thereon;", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=7)
    _p(document, "2.\tTo re-elect the {{ retiring_director_noun }} retiring under the provisions of the Articles of Association of the Company (the Constitution);", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=7)
    _p(document, "3.\tTo re-appoint Messrs. {{ auditor_name }} as auditors of the Company for the ensuing year and to authorise the Directors to fix their remuneration; and", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=7)
    _p(document, "4.\tTo transact any other business for which due notice shall have been given.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=18)
    _p(document, "By Order of the Board,", after=18)
    _p(document, "{{ declarant_name_upper }}", bold=True, after=0)
    _p(document, "Director", after=10)
    _p(document, "Date: {{ notice_date }}", after=9)
    _p(document, "Note: A member of the Company entitled to attend and vote is entitled to appoint a proxy to attend and vote instead of the member. A proxy may but need not be a member of the Company. The instrument appointing a proxy must be deposited at the Company’s Registered Office, {{ meeting_venue_inline }}, not less than forty-eight hours before the time appointed for holding the meeting and at any adjournment thereof.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10)

    _new_page(document)
    _company_header(document)
    _title(document, "MINUTES OF {{ agm_title }}")
    _p(document, "MINUTES OF THE {{ agm_title }} OF THE COMPANY HELD AT {{ meeting_venue_inline }} ON {{ meeting_date }} AT {{ meeting_start_time }}.", bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)
    _p(document, "ATTENDANCE\t:  As per attendance list", after=5)
    _resolution(document, 1, "COMMENCEMENT", "The meeting commenced at {{ meeting_start_time }} with the requisite quorum being present.")
    _resolution(document, 2, "ELECTION OF CHAIRMAN", "{{ chair_with_title }} was elected to chair the meeting.")
    _resolution(document, 3, "NOTICE OF MEETING", "The notice convening the meeting was taken as read.")
    _resolution(document, 4, "AUDITED FINANCIAL STATEMENTS FOR THE FINANCIAL YEAR ENDED {{ financial_year_end_upper }}", "It was resolved that the Audited Financial Statements for the financial year ended {{ financial_year_end }} together with the Reports of the Directors and Auditors thereon be received.")
    _resolution(document, 5, "ROTATION OF {{ retiring_director_noun_upper }}", "That in accordance with Regulation 63 of the Company’s Constitution (Table A in the Fourth Schedule of the Companies Act, 1965 as adopted at that material time), {{ retiring_directors_with_titles }} retired and, being eligible, {{ retiring_offer_verb }} {{ retiring_pronoun }} for re-election. It was resolved that {{ retiring_directors_with_titles }} be re-elected as {{ retiring_director_noun }} of the Company.")
    _resolution(document, 6, "RE-APPOINTMENT OF AUDITORS", "It was resolved that Messrs. {{ auditor_name }} be re-appointed as auditors of the Company for the ensuing year and that the Directors be authorised to fix their remuneration.")
    _resolution(document, 7, "TERMINATION", "There being no other business to discuss, the meeting ended at {{ meeting_end_time }} with a vote of thanks to the Chairman.")
    _p(document, "CONFIRMED CORRECT,", bold=True, before=4, after=14)
    _p(document, "{{ chair_name_upper }}", bold=True, after=0)
    _p(document, "Chairman", after=0)

    _new_page(document)
    _company_header(document)
    _title(document, "ATTENDANCE LIST")
    details = (
        ("Company", "{{ company_name }}"),
        ("Registration No.", "{{ registration_no }}"),
        ("Type of Meeting", "{{ agm_title }}"),
        ("Place", "{{ meeting_venue }}"),
        ("Date & Time", "{{ meeting_date }} at {{ meeting_start_time }}"),
    )
    table = document.add_table(rows=len(details), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(5.0)
    for row, (label, value) in zip(table.rows, details):
        for cell in row.cells:
            _set_cell_margins(cell)
        row.cells[0].text = label
        row.cells[1].text = value
    _p(document, "ATTENDANCE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=4)
    _p(document, "{{ block:attendance_signatures }}", after=0)

    _new_page(document)
    _company_header(document)
    _p(document, "428, Jalan Legenda 26, Legenda Heights, 08000 Sungai Petani, Kedah", align=WD_ALIGN_PARAGRAPH.CENTER, after=9)
    _p(document, "Date: {{ letter_date }}", after=8)
    _p(document, "To:", after=2)
    _p(document, "{{ block:shareholder_addresses }}", after=2)
    _p(document, "{{ salutation }}", before=4, after=8)
    _p(document, "REPORT AND AUDITED FINANCIAL STATEMENTS FOR THE FINANCIAL YEAR ENDED {{ financial_year_end_upper }}", bold=True, after=7)
    _p(document, "We hereby enclose the Reports and Audited Financial Statements for the financial year ended {{ financial_year_end }} pursuant to Section 257 of the Companies Act, 2016 for your attention and safekeeping.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=7)
    _p(document, "Kindly acknowledge receipt.", after=7)
    _p(document, "Thank you.", after=8)
    _p(document, "Yours sincerely,", after=1)
    _p(document, "{{ company_name }}", bold=True, after=16)
    _p(document, "{{ declarant_name_upper }}", bold=True, after=0)
    _p(document, "Director", after=0)

    _new_page(document)
    _company_header(document)
    _title(document, "REPORT AND AUDITED FINANCIAL STATEMENTS FOR THE FINANCIAL YEAR ENDED {{ financial_year_end_upper }}")
    _p(document, "Acknowledged received by:", before=12, after=4)
    _p(document, "{{ block:acknowledgement_signatures }}", after=0)

    document.save(SECTION_90)


def _paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _paragraphs(cell)


def _build_first_agm() -> None:
    if LEGACY_FIRST_AGM.is_file():
        shutil.copy2(LEGACY_FIRST_AGM, FIRST_AGM)
    if not FIRST_AGM.is_file():
        raise FileNotFoundError(f"First-AGM template was not found: {FIRST_AGM}")
    document = Document(FIRST_AGM)
    for paragraph in _paragraphs(document):
        if "(date of incorporation)" not in paragraph.text:
            continue
        for run in paragraph.runs:
            run.text = run.text.replace("(date of incorporation)", "")
    document.save(FIRST_AGM)


def main() -> None:
    TEMPLATES.mkdir(parents=True, exist_ok=True)
    _build_first_agm()
    _build_section_90()
    print(FIRST_AGM)
    print(SECTION_90)


if __name__ == "__main__":
    main()
