from __future__ import annotations

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT_NAME = "Times New Roman"
DEFAULT_SIGNATURE_LINE_SPACE_BEFORE_PT = 34.0


def apply_run_font(run, size: float = 12, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attribute}"), FONT_NAME)


def set_cell_margins(cell, top=80, start=60, bottom=80, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width.twips))
    tc_w.set(qn("w:type"), "dxa")


def configure_layout_table(table, widths: tuple[float, ...]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(Inches(width).twips for width in widths)))
    table_width.set(qn("w:type"), "dxa")
    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for grid_column, width in zip(grid_columns, widths):
        grid_column.set(qn("w:w"), str(Inches(width).twips))
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table, "nil")


def set_table_indent(table, twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(twips))
    indent.set(qn("w:type"), "dxa")


def set_table_borders(table, value: str = "nil") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)


def set_paragraph_border(paragraph, edge: str = "bottom", style: str = "single") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), style)
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "000000")


def set_cell_border(cell, edge: str, style: str = "single", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), style)
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")


def clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_signature_cell(
    cell,
    name: str,
    name_prefix: str = "",
    line_space_before_pt: float = DEFAULT_SIGNATURE_LINE_SPACE_BEFORE_PT,
) -> None:
    clear_cell(cell)
    line = cell.paragraphs[0]
    line.paragraph_format.space_before = Pt(line_space_before_pt)
    line.paragraph_format.space_after = Pt(2)
    line.add_run(" ")
    set_paragraph_border(line)
    name_paragraph = cell.add_paragraph()
    name_paragraph.paragraph_format.space_before = Pt(0)
    name_paragraph.paragraph_format.space_after = Pt(2)
    run = name_paragraph.add_run(f"{name_prefix}{name}")
    apply_run_font(run, 12)
