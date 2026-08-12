from __future__ import annotations

from copy import deepcopy
from decimal import Decimal
from pathlib import Path

from docx import Document

from .models import NewIncorpContext, PersonContext
from .value_utils import (
    address_lines,
    display_name,
    format_date,
    format_decimal,
    format_nric,
    format_percentage,
    normalized_address,
)


DIRECTOR_LEFT = "{{ director_signature_left }}"
DIRECTOR_RIGHT = "{{ director_signature_right }}"


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _replace_markers_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    for marker, value in replacements.items():
        while marker in "".join(run.text for run in paragraph.runs):
            full_text = "".join(run.text for run in paragraph.runs)
            start = full_text.index(marker)
            end = start + len(marker)
            cursor = 0
            first_index = None
            last_index = None
            first_offset = 0
            last_offset = 0
            for index, run in enumerate(paragraph.runs):
                run_end = cursor + len(run.text)
                if first_index is None and start < run_end:
                    first_index = index
                    first_offset = start - cursor
                if end <= run_end:
                    last_index = index
                    last_offset = end - cursor
                    break
                cursor = run_end
            if first_index is None or last_index is None:
                break
            first = paragraph.runs[first_index]
            last = paragraph.runs[last_index]
            if first_index == last_index:
                first.text = first.text[:first_offset] + value + first.text[last_offset:]
            else:
                first.text = first.text[:first_offset] + value
                for index in range(first_index + 1, last_index):
                    paragraph.runs[index].text = ""
                last.text = last.text[last_offset:]


def _iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        seen: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                identity = id(cell._tc)
                if identity in seen:
                    continue
                seen.add(identity)
                yield from _iter_paragraphs(cell)


def _all_story_paragraphs(document: Document):
    yield from _iter_paragraphs(document)
    seen: set[int] = set()
    for section in document.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            identity = id(story._element)
            if identity in seen:
                continue
            seen.add(identity)
            yield from _iter_paragraphs(story)


def _clear_cell(cell) -> None:
    cell.text = ""


def _render_signature_table(table, directors: tuple[PersonContext, ...]) -> None:
    if len(table.rows) != 2 or len(table.columns) != 3:
        raise ValueError("Director signature template must start as a 2x3 table.")
    base_line = deepcopy(table.rows[0]._tr)
    base_name = deepcopy(table.rows[1]._tr)
    pair_count = (len(directors) + 1) // 2
    for _ in range(1, pair_count):
        table._tbl.append(deepcopy(base_line))
        table._tbl.append(deepcopy(base_name))
    for pair in range(pair_count):
        line_row = table.rows[pair * 2]
        name_row = table.rows[pair * 2 + 1]
        left_index = pair * 2
        right_index = left_index + 1
        _replace_markers_in_paragraph(
            name_row.cells[0].paragraphs[0],
            {DIRECTOR_LEFT: directors[left_index].name.upper(), DIRECTOR_RIGHT: directors[left_index].name.upper()},
        )
        if right_index < len(directors):
            _replace_markers_in_paragraph(
                name_row.cells[2].paragraphs[0],
                {DIRECTOR_RIGHT: directors[right_index].name.upper(), DIRECTOR_LEFT: directors[right_index].name.upper()},
            )
        else:
            _clear_cell(line_row.cells[2])
            _clear_cell(name_row.cells[2])
        _clear_cell(line_row.cells[1])
        _clear_cell(name_row.cells[1])


def _render_signature_grids(document: Document, directors: tuple[PersonContext, ...]) -> None:
    signature_tables = [
        table
        for table in document.tables
        if DIRECTOR_LEFT in _table_text(table) or DIRECTOR_RIGHT in _table_text(table)
    ]
    for table in signature_tables:
        _render_signature_table(table, directors)


def _phone(value: str) -> str:
    digits = "".join(character for character in value if character.isdigit())
    if len(digits) in {10, 11}:
        return f"{digits[:3]}-{digits[3:]}"
    return value


def _share_summary(person: PersonContext) -> str:
    shares = format_decimal(person.shares)
    share_class = person.share_class.strip().title() or "Ordinary"
    share_class = share_class.removesuffix(" Shares").removesuffix(" Share")
    noun = "Share" if person.shares == Decimal("1") else "Shares"
    return f"{shares} {share_class} {noun}"


def _common_values(context: NewIncorpContext) -> dict[str, str]:
    registered = address_lines(context.registered_address, 3, 37)
    registered_compact = address_lines(context.registered_address, 2, 52)
    business = address_lines(context.business_address, 3, 37)
    directors_sentence = (
        context.directors[0].name.upper()
        if len(context.directors) == 1
        else ", ".join(person.name.upper() for person in context.directors[:-1])
        + " and "
        + context.directors[-1].name.upper()
    )
    subscriber_rows = "\n".join(
        f"{person.name.upper()}\t{format_decimal(person.shares)}"
        for person in context.members
    )
    return {
        "{{ company_name }}": context.company_name.upper(),
        "{{ registration_no }}": context.registration_no,
        "{{ incorporation_date }}": format_date(context.incorporation_date),
        "{{ incorporation_date_upper }}": format_date(context.incorporation_date, uppercase=True),
        "{{ registered_address }}": normalized_address(context.registered_address),
        "{{ registered_address_line_1 }}": registered[0],
        "{{ registered_address_line_2 }}": registered[1],
        "{{ registered_address_line_3 }}": registered[2],
        "{{ registered_address_compact_line_1 }}": registered_compact[0],
        "{{ registered_address_compact_line_2 }}": registered_compact[1],
        "{{ business_address }}": normalized_address(context.business_address),
        "{{ business_address_upper }}": normalized_address(context.business_address).upper(),
        "{{ business_address_line_1 }}": business[0],
        "{{ business_address_line_2 }}": business[1],
        "{{ business_address_line_3 }}": business[2],
        "{{ director_names_joined }}": " & ".join(person.name.upper() for person in context.directors),
        "{{ director_names_sentence }}": directors_sentence,
        "{{ subscriber_rows }}": subscriber_rows,
        "{{ director1_name }}": context.directors[0].name.upper(),
    }


def _person_values(person: PersonContext, context: NewIncorpContext) -> dict[str, str]:
    residential_upper = tuple(
        line.upper() for line in address_lines(person.residential_address, 3, 34)
    )
    member_title = address_lines(person.residential_address, 3, 34)
    member_compact = address_lines(person.residential_address, 2, 48)
    id_number = (
        format_nric(person.id_number)
        if any(label in person.id_type.upper() for label in ("NRIC", "MYKAD"))
        else person.id_number
    )
    percentage = person.direct_percentage or Decimal("0")
    direct = person.beneficial_owner and percentage >= Decimal("20")
    ordinary = "ORDINARY" in person.share_class.upper()
    return {
        "{{ person_name }}": person.name.upper(),
        "{{ person_nationality_upper }}": person.nationality.upper(),
        "{{ person_race_upper }}": person.race.upper(),
        "{{ person_id }}": id_number,
        "{{ person_dob_upper }}": format_date(person.date_of_birth, uppercase=True),
        "{{ person_phone }}": _phone(person.phone),
        "{{ person_email }}": person.email,
        "{{ person_occupation_upper }}": person.occupation.upper(),
        "{{ residential_address_upper_line_1 }}": residential_upper[0],
        "{{ residential_address_upper_line_2 }}": residential_upper[1],
        "{{ residential_address_upper_line_3 }}": residential_upper[2],
        "{{ member_name }}": person.name.upper(),
        "{{ member_id }}": id_number,
        "{{ member_email }}": person.email,
        "{{ member_phone }}": _phone(person.phone),
        "{{ member_dob }}": format_date(person.date_of_birth),
        "{{ member_nationality_title }}": person.nationality.title(),
        "{{ member_race_title }}": person.race.title(),
        "{{ member_gender_title }}": person.gender.title(),
        "{{ member_occupation_title }}": person.occupation.title(),
        "{{ member_share_summary }}": _share_summary(person),
        "{{ share_fraction }}": f"{format_decimal(person.shares)} of the total {format_decimal(context.total_subscriber_shares)}",
        "{{ direct_percentage }}": format_percentage(person.direct_percentage),
        "{{ becoming_bo_date }}": format_date(person.becoming_bo_date) if person.becoming_bo_date else "",
        "{{ member_address_title_line_1 }}": member_title[0],
        "{{ member_address_title_line_2 }}": member_title[1],
        "{{ member_address_title_line_3 }}": member_title[2],
        "{{ member_address_compact_line_1 }}": member_compact[0],
        "{{ member_address_compact_line_2 }}": member_compact[1],
        "{{ answer_yes_mark }}": "√" if person.beneficial_owner else "",
        "{{ answer_no_mark }}": "" if person.beneficial_owner else "√",
        "{{ individual_type_mark }}": "√" if person.beneficial_owner else "",
        "{{ direct_category_mark }}": "√" if direct else "",
        "{{ criteria_a_mark }}": "√" if direct else "",
        "{{ criteria_b_mark }}": "√" if direct and ordinary else "",
        "{{ criteria_c_mark }}": "√" if person.control_by_other_means else "",
    }


def render_template(
    template_path: Path,
    output_path: Path,
    context: NewIncorpContext,
    person: PersonContext | None = None,
) -> Path:
    document = Document(str(template_path))
    _render_signature_grids(document, context.directors)
    values = _common_values(context)
    if person is not None:
        values.update(_person_values(person, context))
    paragraphs = list(_all_story_paragraphs(document))
    for paragraph in paragraphs:
        _replace_markers_in_paragraph(paragraph, values)
    unresolved = [
        paragraph.text
        for paragraph in _all_story_paragraphs(document)
        if "{{" in paragraph.text or "}}" in paragraph.text
    ]
    if unresolved:
        raise ValueError("Unresolved template markers: " + "; ".join(unresolved[:10]))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
