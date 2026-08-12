from __future__ import annotations

from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document

from .models import NewIncorpContext, ValidationIssue


SAMPLE_VALUES = (
    "HOSAY 3 BAKERY SDN. BHD.",
    "202401051970 (1597813-X)",
    "TAM WEE SEONG",
    "TAN HUI KEE",
    "1181, JALAN PAYA NAHU 1",
)


def _all_text(document: Document) -> str:
    stories: list[str] = []

    def collect(parent) -> None:
        stories.extend(paragraph.text for paragraph in parent.paragraphs)
        seen: set[int] = set()
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    identity = id(cell._tc)
                    if identity in seen:
                        continue
                    seen.add(identity)
                    collect(cell)

    collect(document)
    for section in document.sections:
        collect(section.header)
        collect(section.first_page_header)
        collect(section.even_page_header)
        collect(section.footer)
        collect(section.first_page_footer)
        collect(section.even_page_footer)
    return "\n".join(stories)


def validate_output_docx(
    path: Path,
    context: NewIncorpContext,
    *,
    require_registration_no: bool = True,
) -> list[ValidationIssue]:
    if not path.is_file() or path.stat().st_size == 0:
        return [ValidationIssue("missing_output", f"Generated DOCX was not created: {path}", "output")]
    try:
        with ZipFile(path) as package:
            if "word/document.xml" not in package.namelist():
                raise KeyError("word/document.xml")
        document = Document(str(path))
    except (BadZipFile, KeyError, ValueError) as error:
        return [ValidationIssue("invalid_docx", f"Generated DOCX is invalid: {error}", "output")]
    text = _all_text(document)
    issues: list[ValidationIssue] = []
    if "{{" in text or "}}" in text:
        issues.append(ValidationIssue("unresolved_marker", "Generated output contains unresolved template markers.", "output"))
    if context.company_name.upper() not in text.upper():
        issues.append(ValidationIssue("missing_company_name", "Generated output does not contain the company name.", "output"))
    if require_registration_no and context.registration_no not in text:
        issues.append(ValidationIssue("missing_registration_no", "Generated output does not contain the registration number.", "output"))
    for sample in SAMPLE_VALUES:
        if sample.upper() in text.upper() and sample.upper() not in {
            context.company_name.upper(),
            context.registration_no.upper(),
            *(person.name.upper() for person in context.directors),
            *(person.name.upper() for person in context.members),
        }:
            issues.append(ValidationIssue("sample_value_leaked", f"Template sample value remains: {sample}", "output"))
    return issues
