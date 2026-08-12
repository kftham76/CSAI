"""Build production marker templates from the retained source DOCX copies.

This utility never edits the retained sources. It is intentionally deterministic:
each run starts from originals-before-marker-conversion and writes only the eleven
named production templates in the parent new_incorp template directory.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TEMPLATE_ROOT = Path(__file__).resolve().parent.parent / "templates" / "new_incorp"
SOURCE_ROOT = TEMPLATE_ROOT / "originals-before-marker-conversion"

SOURCE_TO_OUTPUT = {
    "S236(3) Declaration by person before appointment as secretary  (Hosay 3 Bakery).docx": "s236_secretary_declaration_template.docx",
    "0. Beneficial Ownership - Letter to clients on changes (Hosay 3 Bakery).docx": "bo_client_changes_letter_template.docx",
    "1. BO Notice & Reply to Shh  (Individual) - (Hosay 3 Bakery) - TWS.docx": "bo_notice_reply_individual_template.docx",
    "Adopt Policy of BO Reporting (Hosay 3 Bakery).docx": "adopt_bo_policy_template.docx",
    "Disclosure_by_director_-_Tam Wee Seong.docx": "disclosure_director_template.docx",
    "Disclosure_by_member_-_Tam Wee Seong.docx": "disclosure_member_template.docx",
    "DWR - Accounting record kept (Hosay 3 Bakery).docx": "dwr_accounting_records_template.docx",
    "DWR - Appoint Secretary (Hosay 3 Bakery).docx": "dwr_appoint_secretary_template.docx",
    "DWR Authority to Lodge Beneficial Ownership (Hosay 3 Bakery).docx": "dwr_authority_bo_template.docx",
    "DWR_-_FBODM   dated 15 Oct 2024 (Hosay 3 Bakery).docx": "dwr_first_board_meeting_template.docx",
    "Engagement Letter (Hosay 3 Bakery).docx": "engagement_letter_template.docx",
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _roots(document: Document) -> list:
    result = [document.element]
    seen: set[int] = {id(document.element)}
    for section in document.sections:
        for part in (section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer):
            element = part._element
            if id(element) not in seen:
                seen.add(id(element))
                result.append(element)
    return result


def _paragraph_text(paragraph_element) -> str:
    return "".join(paragraph_element.xpath(".//w:t/text()"))


def _set_node_text(node, value: str) -> None:
    node.text = value
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    if value[:1].isspace() or value[-1:].isspace():
        node.set(xml_space, "preserve")
    else:
        node.attrib.pop(xml_space, None)


def _replace_in_paragraph(paragraph_element, old: str, new: str) -> int:
    nodes = paragraph_element.xpath(".//w:t")
    count = 0
    while nodes:
        combined = "".join(node.text or "" for node in nodes)
        start = combined.find(old)
        if start < 0:
            break
        end = start + len(old)
        offsets: list[tuple[int, int]] = []
        cursor = 0
        for node in nodes:
            stop = cursor + len(node.text or "")
            offsets.append((cursor, stop))
            cursor = stop
        first = next(i for i, (_, stop) in enumerate(offsets) if start < stop)
        last = next(i for i, (_, stop) in enumerate(offsets) if end <= stop)
        prefix = (nodes[first].text or "")[: start - offsets[first][0]]
        suffix = (nodes[last].text or "")[end - offsets[last][0] :]
        if first == last:
            _set_node_text(nodes[first], prefix + new + suffix)
        else:
            _set_node_text(nodes[first], prefix + new)
            for index in range(first + 1, last):
                _set_node_text(nodes[index], "")
            _set_node_text(nodes[last], suffix)
        count += 1
    return count


def replace_all(document: Document, old: str, new: str) -> int:
    return sum(
        _replace_in_paragraph(paragraph, old, new)
        for root in _roots(document)
        for paragraph in root.xpath(".//w:p")
    )


def replace_exact(document: Document, old: str, replacements: Iterable[str]) -> None:
    targets = [
        paragraph
        for root in _roots(document)
        for paragraph in root.xpath(".//w:p")
        if _paragraph_text(paragraph).strip() == old.strip()
    ]
    values = list(replacements)
    if len(targets) != len(values):
        raise ValueError(f"Expected {len(values)} exact paragraph(s) for {old!r}, found {len(targets)}")
    for paragraph, value in zip(targets, values):
        _replace_in_paragraph(paragraph, _paragraph_text(paragraph), value)


def _remove_paragraph(paragraph) -> None:
    paragraph._p.getparent().remove(paragraph._p)


def _mark_signature_table(table) -> None:
    table_text = "\n".join(
        cell.text for row in table.rows for cell in row.cells
    )
    if "TAM WEE SEONG" not in table_text or "TAN HUI KEE" not in table_text:
        return
    _replace_in_paragraphs = table._tbl.xpath(".//w:p")
    for paragraph in _replace_in_paragraphs:
        _replace_in_paragraph(paragraph, "TAM WEE SEONG", "{{ director_signature_left }}")
        _replace_in_paragraph(paragraph, "TAN HUI KEE", "{{ director_signature_right }}")


def _mark_all_signature_tables(document: Document) -> None:
    for table in document.tables:
        if len(table.rows) == 2 and len(table.columns) == 3:
            _mark_signature_table(table)


def _signature_exemplar():
    source = Document(str(SOURCE_ROOT / "DWR - Appoint Secretary (Hosay 3 Bakery).docx"))
    heading = next(p for p in source.paragraphs if p.text.strip() == "BOARD OF DIRECTORS")
    return deepcopy(heading._p), deepcopy(source.tables[0]._tbl)


def _replace_sym_with_marker(table, character: str, markers: list[str]) -> None:
    symbols = [node for node in table._tbl.xpath(".//w:sym") if node.get(qn("w:char")) == character]
    if len(symbols) < len(markers):
        raise ValueError(f"Expected at least {len(markers)} symbols {character}, found {len(symbols)}")
    for symbol, marker in zip(symbols, markers):
        run = symbol.getparent()
        for child in list(run):
            if child.tag != qn("w:rPr"):
                run.remove(child)
        text = OxmlElement("w:t")
        text.text = marker
        run.append(text)


def _add_marker_to_cell(cell, marker: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.add_run(marker)


def _table_with_text(document: Document, needle: str):
    for table in document.tables:
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if needle in text:
            return table
    raise ValueError(f"Table containing {needle!r} was not found")


def _common(document: Document) -> None:
    replace_all(document, "HOSAY 3 BAKERY SDN. BHD.", "{{ company_name }}")
    replace_all(document, "202401051970 (1597813-X)", "{{ registration_no }}")


def _build_s236(document: Document) -> None:
    _common(document)
    replace_all(document, "9 DECEMBER 2024", "{{ incorporation_date_upper }}")


def _build_client_letter(document: Document) -> None:
    heading, signature = _signature_exemplar()
    old_table = document.tables[0]._tbl
    old_table.addprevious(heading)
    old_table.addprevious(signature)
    old_table.getparent().remove(old_table)
    _mark_all_signature_tables(document)
    _common(document)
    replace_all(document, "9 December 2024", "{{ incorporation_date }}")
    replace_exact(document, "11 Jalan Seruling 1-A, Taman Seruling,", ["{{ business_address_line_1 }}"])
    replace_exact(document, "08000 Sungai Petani, Kedah", ["{{ business_address_line_2 }}"])


def _build_bo_notice(document: Document) -> None:
    heading, signature = _signature_exemplar()
    paragraphs = document.paragraphs
    line_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "___________________________________")
    anchor = paragraphs[line_index]
    anchor._p.addprevious(heading)
    anchor._p.addprevious(signature)
    for paragraph in paragraphs[line_index : line_index + 3]:
        _remove_paragraph(paragraph)
    _mark_all_signature_tables(document)
    _common(document)
    replace_all(document, "TAM WEE SEONG", "{{ member_name }}")
    replace_all(document, "811229-02-5695", "{{ member_id }}")
    replace_all(document, "kenttam7929@gmail.com", "{{ member_email }}")
    replace_all(document, "017-5717592", "{{ member_phone }}")
    replace_all(document, "29 December 1981", "{{ member_dob }}")
    replace_all(document, "Malaysian", "{{ member_nationality_title }}")
    replace_all(document, "Chinese", "{{ member_race_title }}")
    replace_all(document, "Male", "{{ member_gender_title }}")
    replace_all(document, "Director", "{{ member_occupation_title }}")
    replace_all(document, "1 Ordinary Share", "{{ member_share_summary }}")
    replace_all(document, "1 of the total 2", "{{ share_fraction }}")
    replace_all(document, "50.00", "{{ direct_percentage }}")
    replace_all(document, "9 December 2024", "{{ becoming_bo_date }}")
    replace_exact(document, "428, Jalan Legenda 26, Legenda Heights,", ["{{ registered_address_compact_line_1 }}"])
    replace_exact(
        document,
        "08000 Sungai Petani, Kedah",
        ["{{ registered_address_compact_line_2 }}", "{{ registered_address_line_3 }}"],
    )
    replace_exact(document, "428, Jalan Legenda 26", ["{{ registered_address_line_1 }}"])
    replace_exact(document, "Legenda Heights,", ["{{ registered_address_line_2 }}"])
    replace_exact(document, "No 60 Taman Tunku Husna,", ["{{ member_address_title_line_1 }}"])
    replace_exact(document, "Jalan Tanjung Bendahara,", ["{{ member_address_title_line_2 }}"])
    replace_exact(document, "05300 Alor Setar, Kedah", ["{{ member_address_title_line_3 }}"])
    replace_all(document, " No 60 Taman Tunku Husna, Jalan Tanjung", "{{ member_address_compact_line_1 }}")
    replace_all(document, " Bendahara, 05300 Alor Setar, Kedah", "{{ member_address_compact_line_2 }}")

    answer_table = _table_with_text(document, "I am beneficial owner of the Company")
    _replace_sym_with_marker(answer_table, "F050", ["{{ answer_yes_mark }}"])
    _add_marker_to_cell(answer_table.cell(4, 1), "{{ answer_no_mark }}")
    detail_table = _table_with_text(document, "Details Of The Beneficial Owner")
    _replace_sym_with_marker(detail_table, "F0FE", ["{{ individual_type_mark }}", "{{ direct_category_mark }}"])
    criteria_table = _table_with_text(document, "CATEGORY OF BO")
    _replace_sym_with_marker(
        criteria_table,
        "F050",
        ["{{ criteria_a_mark }}", "{{ criteria_b_mark }}", "{{ criteria_c_mark }}"],
    )


def _build_policy(document: Document) -> None:
    _mark_all_signature_tables(document)
    _common(document)
    replace_all(document, "9 December 2024", "{{ incorporation_date }}")
    replace_all(document, "TAM WEE SEONG & TAN HUI KEE", "{{ director_names_joined }}")


def _build_disclosure(document: Document, role: str) -> None:
    _common(document)
    replace_all(document, "TAM WEE SEONG", "{{ person_name }}")
    replace_all(document, "MALAYSIAN", "{{ person_nationality_upper }}")
    replace_all(document, "CHINESE", "{{ person_race_upper }}")
    replace_all(document, "811229-02-5695", "{{ person_id }}")
    replace_all(document, "29 DECEMBER 1981", "{{ person_dob_upper }}")
    replace_all(document, "017-5717592", "{{ person_phone }}")
    replace_all(document, "kenttam7929@gmail.com", "{{ person_email }}")
    replace_all(document, "DIRECTOR", "{{ person_occupation_upper }}")
    replace_all(document, "NO 60 TAMAN TUNKU HUSNA,", "{{ residential_address_upper_line_1 }}")
    replace_all(document, "JALAN TANJUNG BENDAHARA,", "{{ residential_address_upper_line_2 }}")
    replace_all(document, "05300 ALOR SETAR, KEDAH", "{{ residential_address_upper_line_3 }}")
    replace_all(
        document,
        "11 JALAN SERULING 1-A, TAMAN SERULING, 08000 SUNGAI PETANI, KEDAH",
        "{{ business_address_upper }}",
    )


def _build_accounting(document: Document) -> None:
    _mark_all_signature_tables(document)
    _common(document)
    replace_all(
        document,
        "11 Jalan Seruling 1-A, Taman Seruling, 08000 Sungai Petani, Kedah",
        "{{ business_address }}",
    )
    replace_all(document, "11 Jalan Seruling 1-A, Taman Seruling,", "{{ business_address_line_1 }}")
    replace_all(document, "08000 Sungai Petani, Kedah", "{{ business_address_line_2 }}")


def _build_appoint(document: Document) -> None:
    _mark_all_signature_tables(document)
    _common(document)
    replace_all(document, "9 December 2024", "{{ incorporation_date }}")


def _build_authority(document: Document) -> None:
    _mark_all_signature_tables(document)
    _common(document)


def _build_first_board(document: Document) -> None:
    _mark_all_signature_tables(document)
    _common(document)
    replace_all(document, "9 December 2024", "{{ incorporation_date }}")
    replace_all(document, "TAM WEE SEONG and TAN HUI KEE", "{{ director_names_sentence }}")
    replace_all(
        document,
        "428, Jalan Legenda 26, Legenda Heights, 08000 Sungai Petani, Kedah",
        "{{ registered_address }}",
    )
    replace_exact(document, "TAM WEE SEONG     1", ["{{ subscriber_rows }}"])
    replace_exact(document, "TAN HUI KEE     1", [""])


def _build_engagement(document: Document) -> None:
    _common(document)
    replace_all(document, "9 December 2024", "{{ incorporation_date }}")
    replace_exact(document, "1181, Jalan Paya Nahu 1", ["{{ business_address_line_1 }}"])
    replace_exact(document, "Kampung Raja", ["{{ business_address_line_2 }}"])
    replace_exact(document, "08000 Sungai Petani, Kedah", ["{{ business_address_line_3 }}"])
    replace_all(document, "TAM WEE SEONG", "{{ director1_name }}")


BUILDERS = {
    "s236_secretary_declaration_template.docx": _build_s236,
    "bo_client_changes_letter_template.docx": _build_client_letter,
    "bo_notice_reply_individual_template.docx": _build_bo_notice,
    "adopt_bo_policy_template.docx": _build_policy,
    "disclosure_director_template.docx": lambda document: _build_disclosure(document, "director"),
    "disclosure_member_template.docx": lambda document: _build_disclosure(document, "member"),
    "dwr_accounting_records_template.docx": _build_accounting,
    "dwr_appoint_secretary_template.docx": _build_appoint,
    "dwr_authority_bo_template.docx": _build_authority,
    "dwr_first_board_meeting_template.docx": _build_first_board,
    "engagement_letter_template.docx": _build_engagement,
}


def build_templates() -> list[Path]:
    outputs: list[Path] = []
    for source_name, output_name in SOURCE_TO_OUTPUT.items():
        source = SOURCE_ROOT / source_name
        if not source.is_file():
            raise FileNotFoundError(source)
        document = Document(str(source))
        BUILDERS[output_name](document)
        output = TEMPLATE_ROOT / output_name
        document.save(str(output))
        outputs.append(output)
    return outputs


def main() -> int:
    for output in build_templates():
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
