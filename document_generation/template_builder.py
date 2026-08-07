from __future__ import annotations

import argparse
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from .docx_helpers import FONT_NAME, apply_run_font, set_paragraph_border


HEADER_RAISE = Cm(1)
STANDARD_TOP_MARGIN = Inches(0.5)
TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
DEFAULT_REFERENCE_DOCX = TEMPLATE_DIR / "agm_approve_accounts_template.docx"
DEFAULT_OUTPUT_DOCX = TEMPLATE_DIR / "agm_approve_accounts_template.docx"


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for style_name in ("List Number",):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = FONT_NAME
            style.font.size = Pt(12)
    if "AGM Block Marker" not in document.styles:
        marker = document.styles.add_style("AGM Block Marker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        marker = document.styles["AGM Block Marker"]
    marker.font.name = FONT_NAME
    marker.font.size = Pt(9)
    marker.font.italic = True


def _set_tab(paragraph, position: float = 0.5) -> None:
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(position))


def _add_header(
    document: Document,
    include_address: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(
        30 if include_address else max(0, 30 - HEADER_RAISE.pt)
    )
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("{{ company_name }}")
    apply_run_font(run, 14, True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Registration No : {{ registration_no }}")
    apply_run_font(run, 11)
    if include_address:
        for line in (
            "428, Jalan Legenda 26",
            "Legenda Heights,",
            "08000 Sungai Petani, Kedah",
        ):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            apply_run_font(paragraph.add_run(line), 11)
    else:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # The DWR/MWR company block is raised by 1 cm at section level. Add
        # the same distance below it so the title and body stay in place.
        paragraph.paragraph_format.space_after = Pt(6 + HEADER_RAISE.pt)
        apply_run_font(paragraph.add_run("(Incorporated in Malaysia)"), 11)


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    apply_run_font(paragraph.add_run(text), 14, True)


def _add_statutory_clause(document: Document, marker: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    apply_run_font(paragraph.add_run(marker), 12)
    set_paragraph_border(paragraph)


def _add_resolution(document: Document, number: str, heading: str, body: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_tab(paragraph)
    apply_run_font(paragraph.add_run(f"{number}\t"), 12)
    apply_run_font(paragraph.add_run(heading), 12, True)
    body_paragraph = document.add_paragraph()
    body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_paragraph.paragraph_format.left_indent = Inches(0.5)
    body_paragraph.paragraph_format.space_after = Pt(2)
    apply_run_font(body_paragraph.add_run(body), 12)


def _add_block_marker(document: Document, marker: str) -> None:
    paragraph = document.add_paragraph(style="AGM Block Marker")
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    apply_run_font(paragraph.add_run(marker), 9)


def _add_signature_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(30)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.right_indent = Inches(4.3)
    paragraph.add_run(" ")
    set_paragraph_border(paragraph)


def _apply_numbering(paragraph, num_id: int = 1) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(number)


def _build_template_from_code(reference_docx: Path, output_docx: Path) -> Path:
    document = Document(str(reference_docx))
    _clear_body(document)
    _configure_styles(document)
    for section in document.sections:
        section.top_margin = STANDARD_TOP_MARGIN

    _add_header(document)
    _add_title(document, "DIRECTORS’ WRITTEN RESOLUTION")
    _add_statutory_clause(document, "{{ dwr_clause }}")
    paragraph = document.add_paragraph("It is hereby resolved:-")
    apply_run_font(paragraph.runs[0], 12)

    _add_resolution(
        document,
        "1.",
        "DIRECTORS’ REPORT",
        "That the Directors’ Report, a copy of which is attached hereto, be approved and adopted as the Directors’ Report required pursuant to Section 252(2) of the Companies Act, 2016 to be made and attached to the Company’s Audited Financial Statements for the financial year ended {{ financial_year_end }}.",
    )
    _add_resolution(
        document,
        "2.",
        "AUDITED FINANCIAL STATEMENTS",
        "That the Balance Sheet as at {{ financial_year_end }} and Income Statement for the financial year ended {{ financial_year_end }} which have been made out in accordance with the applicable approved accounting standards be and are hereby approved.",
    )
    _add_resolution(
        document,
        "3.",
        "SIGNING OF DIRECTORS’ REPORT AND STATEMENT",
        "That {{ statement_signers_with_titles }} {{ statement_signer_authority }} hereby authorised to sign for and on behalf of the Board, pursuant to Sections 252(2)(b), 251(2) and 251(3) of the Companies Act, 2016, the Directors’ Report referred to above and the Directors’ Statement accompanying the Audited Financial Statements for the financial year ended {{ financial_year_end }}.",
    )
    _add_resolution(
        document,
        "4.",
        "SIGNING OF STATUTORY DECLARATION",
        "That {{ declarant_with_title }} be and is hereby authorised to sign, pursuant to Section 251(1)(b) of the Companies Act, 2016, the Statutory Declaration accompanying the Audited Financial Statements for the financial year ended {{ financial_year_end }}.",
    )
    _add_block_marker(document, "{{ block:director_fee_section }}")
    _add_resolution(
        document,
        "{{ circulation_section_number }}.",
        "CIRCULATION OF REPORTS AND FINANCIAL STATEMENTS",
        "That the Company Secretary be and is hereby authorised to circulate the Reports and Audited Financial Statements of the Company for the financial year ended {{ financial_year_end }} to the members pursuant to Section 257 of the Companies Act, 2016 on {{ circulation_date }}.",
    )
    _add_resolution(
        document,
        "{{ filing_section_number }}.",
        "AUTHORITY TO FILE AUDITED FINANCIAL STATEMENTS",
        "That the Company Secretary be authorised to file the Audited Financial Statements for the financial year ended {{ financial_year_end }} with the Companies Commission of Malaysia within 30 days pursuant to Section 259(1) of the Companies Act, 2016.",
    )
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    apply_run_font(heading.add_run("{{ director_signature_heading }}"), 12, True)
    _add_block_marker(document, "{{ block:director_signatures }}")
    paragraph = document.add_paragraph("Date: {{ board_approval_date }}")
    apply_run_font(paragraph.runs[0], 12)

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = STANDARD_TOP_MARGIN
    _add_header(document)
    _add_title(document, "MEMBERS’ WRITTEN RESOLUTION")
    _add_statutory_clause(document, "{{ mwr_clause }}")
    paragraph = document.add_paragraph("It was resolved:-")
    apply_run_font(paragraph.runs[0], 12)
    _add_resolution(
        document,
        "1.",
        "AUDITED FINANCIAL STATEMENTS FOR THE FINANCIAL YEAR ENDED {{ financial_year_end_upper }}",
        "It was duly acknowledged and noted that the Audited Financial Statements for the financial year ended {{ financial_year_end }}, together with the Reports of the Directors and Auditors thereon, were circulated by the Company Secretary on {{ circulation_date }} to all members pursuant to Section 257 of the Companies Act, 2016.",
    )
    _add_resolution(
        document,
        "2.",
        "RE-APPOINTMENT OF AUDITORS",
        "That Messrs. {{ auditor_name }} be re-appointed as auditors of the Company for the ensuing year and that the Directors be authorised to fix their remuneration.",
    )
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(14)
    apply_run_font(heading.add_run("{{ member_signature_heading }}"), 12, True)
    _add_block_marker(document, "{{ block:member_signatures }}")
    paragraph = document.add_paragraph("Date: {{ circulation_date }}")
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph(
        "Note: Statement informing members pursuant to Section 301 of the Companies Act, 2016:-"
    )
    paragraph.paragraph_format.space_before = Pt(8)
    apply_run_font(paragraph.runs[0], 12)
    notes = (
        "If you agree to the resolution(s) proposed herein, please sign in the space indicated above your name.",
        "Please return the duly signed resolution(s) to the Registered Office at 428, Jalan Legenda 26, Legenda Heights, 08000 Sungai Petani, Kedah by hand, post, courier or registered post.",
        "If the resolution(s) are not agreed and signed by the required majority of members and returned to the Company by {{ lapse_date }}, the resolution(s) shall lapse and be invalid.",
        "Any agreement given for the above shall not be revoked.",
    )
    for text in notes:
        paragraph = document.add_paragraph()
        _apply_numbering(paragraph)
        paragraph.paragraph_format.space_after = Pt(0)
        apply_run_font(paragraph.add_run(text), 12)

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = STANDARD_TOP_MARGIN
    _add_header(document, include_address=True)
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(8)
    rule.add_run(" ")
    set_paragraph_border(rule)
    paragraph = document.add_paragraph("Date: {{ circulation_date }}")
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    run = paragraph.add_run("{{ shareholder_heading }}")
    apply_run_font(run, 12)
    run.underline = True
    _add_block_marker(document, "{{ block:shareholder_addresses }}")
    paragraph = document.add_paragraph("{{ salutation }}")
    paragraph.paragraph_format.space_before = Pt(8)
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    run = paragraph.add_run(
        "Report and Audited Financial Statements for the financial year ended {{ financial_year_end }}"
    )
    apply_run_font(run, 12, True)
    run.underline = True
    paragraph = document.add_paragraph(
        "We enclose herewith the Audited Financial Statements for the financial year ended {{ financial_year_end }} of {{ company_name }} for your kind attention and safekeeping."
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(8)
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph("Thank you.")
    paragraph.paragraph_format.space_before = Pt(8)
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph("Yours sincerely,")
    paragraph.paragraph_format.space_before = Pt(8)
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph("{{ company_name }}")
    apply_run_font(paragraph.runs[0], 12, True)
    _add_signature_rule(document)
    paragraph = document.add_paragraph("{{ declarant_name_upper }}")
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph("Director")
    apply_run_font(paragraph.runs[0], 12)
    paragraph = document.add_paragraph(
        "Kindly acknowledge receipt of the abovementioned Audited Report."
    )
    paragraph.paragraph_format.space_before = Pt(24)
    apply_run_font(paragraph.runs[0], 12)
    _add_block_marker(document, "{{ block:acknowledgement_signatures }}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return output_docx


def build_template(
    reference_docx: Path = DEFAULT_REFERENCE_DOCX,
    output_docx: Path = DEFAULT_OUTPUT_DOCX,
) -> Path:
    """Clone the approved AGM template without reconstructing its OOXML."""
    reference_docx = Path(reference_docx).resolve()
    output_docx = Path(output_docx).resolve()
    if not reference_docx.is_file():
        raise FileNotFoundError(f"Reference template was not found: {reference_docx}")
    if reference_docx == output_docx:
        raise ValueError("Reference and output template paths must be different.")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    handle = tempfile.NamedTemporaryFile(
        prefix="agm-template-",
        suffix=".docx",
        dir=output_docx.parent,
        delete=False,
    )
    temporary_path = Path(handle.name)
    handle.close()
    try:
        shutil.copy2(reference_docx, temporary_path)
        temporary_path.replace(output_docx)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise
    return output_docx


def main() -> int:
    parser = argparse.ArgumentParser(description="Clone the approved AGM template.")
    parser.add_argument(
        "reference_docx",
        type=Path,
        nargs="?",
        default=DEFAULT_REFERENCE_DOCX,
    )
    parser.add_argument(
        "output_docx",
        type=Path,
        nargs="?",
        default=DEFAULT_OUTPUT_DOCX,
    )
    args = parser.parse_args()
    build_template(args.reference_docx, args.output_docx)
    print(args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
