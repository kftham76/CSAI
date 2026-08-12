from __future__ import annotations

import tempfile
import unittest
from datetime import date, timedelta
from decimal import Decimal
from pathlib import Path

from docx import Document

from document_generation.context_builder import build_document_context
from document_generation.data_provider import CsaiDataProvider
from document_generation.generator import (
    DEFAULT_TEMPLATE,
    IRREGULAR_YEAR_TEMPLATE,
    generate_documents,
    resolve_template_path,
)
from document_generation.output_validation import validate_output_docx
from document_generation.text_utils import normalize_address


class FakeProvider:
    def __init__(
        self,
        member_count: int = 3,
        fee: str = "100.00",
        declarant: str = "PERSON 1",
        financial_year_start: str = "2025-01-01",
        financial_year_end: str = "2025-12-31",
    ) -> None:
        self.member_count = member_count
        company = {
            "Company Name": "TEST COMPANY SDN. BHD.",
            "Reg No": "202601000001 (1700001-A)",
            "Folder": "Test Company",
        }
        for index in range(1, member_count + 1):
            name = f"PERSON {index}"
            if index == member_count and member_count == 6:
                name = "A VERY LONG CORPORATE MEMBER NAME HOLDINGS SDN. BHD."
            company[f"Director{index} Name"] = name
            company[f"Director{index} Gender"] = "MALE" if index % 2 else "FEMALE"
            company[f"Director{index} Residential Address"] = f"Director address {index}"
            company[f"Member{index} Name"] = name
            company[f"Member{index} Gender"] = "MALE" if index % 2 else "FEMALE"
            company[f"Member{index} Address"] = (
                f"No. {index}, A deliberately long residential address, "
                "Taman Example, 08000 Sungai Petani, Kedah, Malaysia"
            )
            company[f"Member{index} Shares"] = "1"
            company[f"Member{index} Type"] = "Corporate" if index == member_count else "Individual"
        self.company = company
        self.financial = {
            "Company's current financial year start date": financial_year_start,
            "Company's current financial year end date": financial_year_end,
            "Date of financial statements approved by Board of Directors": "2026-06-15",
            "Date of circulation of financial statements and reports to members": "2026-06-30",
            "Date of Statutory Declaration": "2026-06-15",
            "Statutory Declaration - Name of director who made declaration": declarant,
            "Number of directors signing Statement by Directors": "2" if member_count > 1 else "1",
            "Name of first director who signed Statement by Directors": "PERSON 1",
            "Name of second director who signed Statement by Directors": "PERSON 2" if member_count > 1 else "",
            "Name of audit firm": "TEST & CO.",
            "Director's remuneration - Fees (Current Financial Year)": fee,
        }
        self.constitution = {
            "Reg No": company["Reg No"],
            "DIRECTOR WRITTEN RESOLUTION (DWR Statutory)": (
                "DIRECTORS' WRITTEN RESOLUTION PASSED PURSUANT TO PARAGRAPH 15 "
                "OF THE THIRD SCHEDULE OF THE COMPANIES ACT, 2016"
            ),
            "MEMBER WRITTEN RESOLUTION (MWR Statutory)": (
                "MEMBERS' WRITTEN RESOLUTION MADE PURSUANT TO SECTION 290(1)(a) OF THE COMPANIES ACT, 2016"
            ),
        }
        self.auditor = {
            "Reg No": company["Reg No"],
            "Auditor Name": "TEST & CO.",
        }

    def get_company(self, company_name: str) -> list[dict]:
        return [self.company]

    def get_financial(self, company_name: str) -> list[dict]:
        return [self.financial]

    def get_constitution(self, company_name: str) -> list[dict]:
        return [self.constitution]

    def get_auditor(self, company_name: str) -> list[dict]:
        return [self.auditor]


class ContextTests(unittest.TestCase):
    def test_address_removes_only_trailing_country(self) -> None:
        normalized = normalize_address(
            "No. 1, Malaysia Park, 08000 Sungai Petani, Kedah, MALAYSIA"
        )
        self.assertIn("Malaysia Park", normalized)
        self.assertFalse(normalized.rstrip().endswith("Malaysia"))
        self.assertFalse(normalized.rstrip().endswith(","))

    def test_fee_rounding_uses_final_row_remainder(self) -> None:
        result = build_document_context("test company", FakeProvider())
        self.assertTrue(result.valid, result.issues)
        amounts = [item.amount for item in result.context.fee_allocations]
        self.assertEqual(amounts, [Decimal("33.33"), Decimal("33.33"), Decimal("33.34")])
        self.assertEqual(sum(amounts), Decimal("100.00"))

    def test_missing_declarant_blocks_generation(self) -> None:
        result = build_document_context("test company", FakeProvider(declarant=""))
        self.assertFalse(result.valid)
        self.assertIn("missing_fs_field", {issue.code for issue in result.issues})

    def test_numeric_person_order_is_preserved(self) -> None:
        result = build_document_context("test company", FakeProvider(member_count=6, fee="0"))
        self.assertTrue(result.valid, result.issues)
        self.assertEqual([person.index for person in result.context.members], [1, 2, 3, 4, 5, 6])

    def test_missing_and_invalid_financial_year_start_are_rejected(self) -> None:
        for value in ("", "not-a-date"):
            with self.subTest(value=value):
                result = build_document_context(
                    "test company",
                    FakeProvider(financial_year_start=value),
                )
                self.assertFalse(result.valid)
                self.assertTrue(
                    any(
                        issue.code == "invalid_fs_date"
                        and "financial year start date" in issue.message
                        for issue in result.issues
                    ),
                    result.issues,
                )

    def test_reversed_financial_year_range_is_rejected(self) -> None:
        result = build_document_context(
            "test company",
            FakeProvider(
                financial_year_start="2026-01-01",
                financial_year_end="2025-12-31",
            ),
        )
        self.assertFalse(result.valid)
        self.assertIn(
            "invalid_financial_year_range",
            {issue.code for issue in result.issues},
        )


class TemplateSelectionTests(unittest.TestCase):
    def test_selection_boundaries(self) -> None:
        financial_year_end = date(2025, 12, 31)
        cases = (
            (363, IRREGULAR_YEAR_TEMPLATE),
            (364, DEFAULT_TEMPLATE),
            (367, DEFAULT_TEMPLATE),
            (368, IRREGULAR_YEAR_TEMPLATE),
        )
        for span_days, expected in cases:
            with self.subTest(span_days=span_days):
                start = financial_year_end - timedelta(days=span_days)
                result = build_document_context(
                    "test company",
                    FakeProvider(
                        financial_year_start=start.isoformat(),
                        financial_year_end=financial_year_end.isoformat(),
                    ),
                )
                self.assertTrue(result.valid, result.issues)
                self.assertEqual(resolve_template_path(result.context), expected)

    def test_explicit_template_overrides_automatic_selection(self) -> None:
        result = build_document_context(
            "test company",
            FakeProvider(financial_year_start="2025-04-03"),
        )
        self.assertTrue(result.valid, result.issues)
        override = Path(r"C:\custom\agm-template.docx")
        self.assertEqual(resolve_template_path(result.context, override), override)

    def test_dry_run_preview_reports_period_and_template(self) -> None:
        results = generate_documents(
            ["test company"],
            dry_run=True,
            provider=FakeProvider(financial_year_start="2025-04-03"),
        )
        self.assertEqual(results[0].status, "valid", results[0].issues)
        self.assertEqual(results[0].preview["financial_year_start"], "2025-04-03")
        self.assertEqual(results[0].preview["financial_year_span_days"], 272)
        self.assertEqual(
            results[0].preview["selected_template_path"],
            str(IRREGULAR_YEAR_TEMPLATE),
        )


class RenderingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_directory = tempfile.TemporaryDirectory()

    def _render(self, count: int, fee: str = "0") -> Path:
        output_dir = Path(self.temp_directory.name) / f"members-{count}-fee-{fee}"
        results = generate_documents(
            ["test company"],
            template_path=DEFAULT_TEMPLATE,
            output_dir=output_dir,
            provider=FakeProvider(member_count=count, fee=fee),
        )
        self.assertEqual(results[0].status, "generated", results[0].issues)
        return results[0].output_path

    def tearDown(self) -> None:
        self.temp_directory.cleanup()

    def test_one_through_six_people_render_without_markers(self) -> None:
        for count in range(1, 7):
            with self.subTest(count=count):
                output = self._render(count)
                self.assertEqual(validate_output_docx(output), [])
                document = Document(str(output))
                text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                self.assertIn("Dear Sir," if count == 1 else "Dear Sirs,", text)
                acknowledgement = [
                    table
                    for table in document.tables
                    if any("Name:" in cell.text for row in table.rows for cell in row.cells)
                ]
                self.assertEqual(len(acknowledgement), 1)
                self.assertEqual(len(acknowledgement[0].rows), (count + 1) // 2)

    def test_zero_fee_removes_section_and_renumbers(self) -> None:
        output = self._render(2, fee="0")
        document = Document(str(output))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertNotIn("PAYMENT OF DIRECTORS", text)
        self.assertIn("5.\tCIRCULATION OF REPORTS", text)
        self.assertIn("6.\tAUTHORITY TO ", text)
        self.assertIn("AUDITED FINANCIAL STATEMENTS", text)

    def test_positive_fee_section_totals_exactly(self) -> None:
        output = self._render(3, fee="100")
        document = Document(str(output))
        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        self.assertIn("PAYMENT OF DIRECTORS", all_text)
        self.assertIn("RM33.34", all_text)
        self.assertIn("RM100.00", all_text)


class TemplateLayoutTests(unittest.TestCase):
    def test_template_contains_only_one_dwr_authority_resolution(self) -> None:
        document = Document(str(DEFAULT_TEMPLATE))
        authority_headings = [
            paragraph
            for paragraph in document.paragraphs
            if "AUTHORITY TO FILE AUDITED FINANCIAL STATEMENTS" in paragraph.text
            or "AUTHORITY TO LODGE AUDITED FINANCIAL STATEMENTS" in paragraph.text
        ]
        self.assertEqual(len(authority_headings), 1)

    def test_all_three_company_blocks_and_two_incorporation_lines_are_present(self) -> None:
        document = Document(str(DEFAULT_TEMPLATE))
        company_paragraphs = [
            paragraph
            for index, paragraph in enumerate(document.paragraphs[:-1])
            if paragraph.text == "{{ company_name }}"
            and document.paragraphs[index + 1].text.startswith("Registration No :")
        ]
        incorporated_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "(Incorporated in Malaysia)"
        ]
        self.assertEqual(len(company_paragraphs), 3)
        self.assertEqual(len(incorporated_paragraphs), 2)

    def test_irregular_template_contains_start_and_end_markers(self) -> None:
        document = Document(str(IRREGULAR_YEAR_TEMPLATE))
        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if "being the first financial year" in paragraph.text
        ]
        self.assertEqual(len(paragraphs), 1)
        self.assertIn("{{ financial_year_start }}", paragraphs[0])
        self.assertEqual(paragraphs[0].count("{{ financial_year_end }}"), 2)
        self.assertNotIn("FS.db", paragraphs[0])
        with tempfile.TemporaryDirectory() as directory:
            result = generate_documents(
                ["test company"],
                output_dir=Path(directory),
                provider=FakeProvider(financial_year_start="2025-04-03"),
            )[0]
            self.assertEqual(result.status, "generated", result.issues)
            generated = Document(str(result.output_path))
            generated_text = "\n".join(
                paragraph.text for paragraph in generated.paragraphs
            )
            self.assertNotIn("date of incorporation", generated_text)


class LiveDatabaseTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.provider = CsaiDataProvider()

    def test_cs_suria_is_valid(self) -> None:
        result = build_document_context("CS SURIA SDN. BHD.", self.provider)
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(len(result.context.members), 4)

    def test_cy_global_fee_split(self) -> None:
        result = build_document_context("CY GLOBAL INDUSTRIES SDN. BHD.", self.provider)
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(
            [allocation.amount for allocation in result.context.fee_allocations],
            [Decimal("84000.00"), Decimal("84000.00"), Decimal("42000.00")],
        )


if __name__ == "__main__":
    unittest.main()
