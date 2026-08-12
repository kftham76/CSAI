from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from document_generation.context_builder import build_document_context
from document_generation.generator import (
    DEFAULT_TEMPLATE,
    FIRST_AGM_TEMPLATE,
    SECTION_90_TEMPLATE,
    generate_documents,
    resolve_template_path,
)
from document_generation.models import Person
from document_generation.output_validation import validate_output_docx
from document_generation.rotation_reader import (
    discover_rotation_workbook,
    read_retiring_directors,
)
from document_generation.template_selection import (
    TEMPLATE_SECTION_90,
    TEMPLATE_STANDARD,
    TEMPLATE_STANDARD_OR_FIRST_AGM,
    classify_dwr,
)

from document_generation.tests.agm.test_generator import FakeProvider


PARAGRAPH_15 = (
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO PARAGRAPH 15 "
    "OF THE THIRD SCHEDULE OF THE COMPANIES ACT, 2016"
)
REGULATION_90 = (
    "DIRECTORS� WRITTEN RESOLUTION PASSED PURSUANT TO REGULATION 90 "
    "(TABLE A) OF THE COMPANY�S ARTICLES OF ASSOCIATION (�THE CONSTITUTION�)"
)


STANDARD_OR_FIRST_AGM_DWRS = (
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 23(b) "
    "OF THE COMPANY’S CONSTITUTION",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 34 "
    "OF THE COMPANY’S CONSTITUTION",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 36 "
    "OF THE COMPANY’S CONSTITUTION",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 37 "
    "OF THE COMPANY’S CONSTITUTION",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO PARAGRAPH 15 "
    "OF THE THIRD SCHEDULE OF THE COMPANIES ACT, 2016",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO REGULATION 34 "
    "OF THE COMPANY’S CONSTITUTION",
)
SECTION_90_DWRS = (
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT ARTICLE 5 OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (“THE CONSTITUTION”)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 3(d) OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (“THE CONSTITUTION”)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 5 OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (“THE CONSTITUTION”)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 72 OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (COMPANY’S CONSTITUTION)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 77 OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (COMPANY’S CONSTITUTION)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 9 OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (“THE CONSTITUTION”)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 95 OF THE "
    "COMPANY’S ARTICLES OF ASSOCIATION (“THE CONSTITUTION”)",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO CLAUSE 53 "
    "OF THE COMPANY’S CONSTITUTION",
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO REGULATION 90 "
    "(TABLE A) OF THE COMPANY’S ARTICLES OF ASSOCIATION (“THE CONSTITUTION”)",
)


def section90_provider(count: int = 3) -> FakeProvider:
    provider = FakeProvider(
        member_count=count,
        fee="not-a-number",
        financial_year_start="2025-04-03",
    )
    provider.constitution["DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"] = REGULATION_90
    provider.constitution["MEMBER WRITTEN RESOLUTION (MWR Statutory)"] = ""
    return provider


def write_rotation(
    clients_root: Path,
    provider: FakeProvider,
    checked_names: list[str],
    *,
    year: int = 2025,
    circulation_subfolder: bool = True,
    filename: str = "Retirement Rotation.xlsx",
) -> Path:
    base = clients_root / provider.company["Folder"] / "AGM"
    if circulation_subfolder:
        base = base / "2026"
    base.mkdir(parents=True, exist_ok=True)
    path = base / filename
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["No.", "Director", year])
    for index in range(1, provider.member_count + 1):
        name = provider.company[f"Director{index} Name"]
        sheet.append([index, name, "✓" if name in checked_names else ""])
    workbook.save(path)
    workbook.close()
    return path


class DwrClassifierTests(unittest.TestCase):
    def test_classifies_all_fifteen_workbook_dwr_strings(self) -> None:
        for value in STANDARD_OR_FIRST_AGM_DWRS:
            with self.subTest(value=value):
                self.assertEqual(
                    classify_dwr(value),
                    TEMPLATE_STANDARD_OR_FIRST_AGM,
                )
        for value in SECTION_90_DWRS:
            with self.subTest(value=value):
                self.assertEqual(classify_dwr(value), TEMPLATE_SECTION_90)

    def test_classifies_user_supplied_and_encoding_variants(self) -> None:
        variants = (
            REGULATION_90,
            "Directors' Written Resolution pursuant to Regulation 90, Table A, "
            "of the Company's Articles of Association (Companies Constitution)",
            "DIRECTORS’ WRITTEN RESOLUTION REGULATION 90 (TABLE A) "
            "ARTICLES OF ASSOCIATION",
        )
        for value in variants:
            with self.subTest(value=value):
                self.assertEqual(classify_dwr(value), TEMPLATE_SECTION_90)

    def test_article_9_does_not_match_article_95_by_prefix(self) -> None:
        article_9 = SECTION_90_DWRS[5]
        article_95 = SECTION_90_DWRS[6]
        self.assertNotEqual(article_9, article_95)
        self.assertEqual(classify_dwr(article_9), TEMPLATE_SECTION_90)
        self.assertEqual(classify_dwr(article_95), TEMPLATE_SECTION_90)

    def test_standard_authorities_require_company_constitution_source(self) -> None:
        variants = (
            "DIRECTORS' WRITTEN RESOLUTION PASSED PURSUANT TO ARTICLE 34 "
            "OF THE COMPANY'S ARTICLES OF ASSOCIATION (THE CONSTITUTION)",
            "DIRECTORS' WRITTEN RESOLUTION PASSED PURSUANT TO REGULATION 34 "
            "OF THE COMPANY'S ARTICLES OF ASSOCIATION (THE CONSTITUTION)",
        )
        for value in variants:
            with self.subTest(value=value):
                self.assertEqual(classify_dwr(value), TEMPLATE_STANDARD)

    def test_unknown_dwr_falls_back_to_template_1_even_for_irregular_period(self) -> None:
        provider = FakeProvider(financial_year_start="2025-04-03")
        provider.constitution["DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"] = "UNKNOWN CLAUSE"
        result = build_document_context("test company", provider)
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(result.context.template_family, TEMPLATE_STANDARD)
        self.assertEqual(resolve_template_path(result.context), DEFAULT_TEMPLATE)

    def test_all_standard_authorities_use_period_rule(self) -> None:
        cases = (
            ("2025-01-02", FIRST_AGM_TEMPLATE),  # 363 days
            ("2025-01-01", DEFAULT_TEMPLATE),  # 364 days
            ("2024-12-29", DEFAULT_TEMPLATE),  # 367 days
            ("2024-12-28", FIRST_AGM_TEMPLATE),  # 368 days
        )
        for dwr in STANDARD_OR_FIRST_AGM_DWRS:
            for start, expected in cases:
                with self.subTest(dwr=dwr, start=start):
                    provider = FakeProvider(
                        financial_year_start=start,
                        financial_year_end="2025-12-31",
                    )
                    provider.constitution[
                        "DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"
                    ] = dwr
                    result = build_document_context("test company", provider)
                    self.assertTrue(result.valid, result.issues)
                    self.assertEqual(
                        result.context.template_family,
                        TEMPLATE_STANDARD_OR_FIRST_AGM,
                    )
                    self.assertEqual(resolve_template_path(result.context), expected)

    def test_all_section_90_authorities_use_full_template_3_context(self) -> None:
        for dwr in SECTION_90_DWRS:
            with self.subTest(dwr=dwr):
                provider = section90_provider()
                provider.constitution[
                    "DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"
                ] = dwr
                result = build_document_context("test company", provider)
                self.assertTrue(result.valid, result.issues)
                self.assertEqual(result.context.template_family, TEMPLATE_SECTION_90)
                self.assertEqual(resolve_template_path(result.context), SECTION_90_TEMPLATE)
                self.assertEqual(result.context.fee_allocations, [])


class RotationWorkbookTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.provider = section90_provider(3)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_prefers_circulation_year_subfolder_and_reads_multiple_names(self) -> None:
        root_path = write_rotation(
            self.root,
            self.provider,
            ["PERSON 3"],
            circulation_subfolder=False,
        )
        preferred = write_rotation(
            self.root,
            self.provider,
            ["PERSON 1", "PERSON 2"],
        )
        discovered, issues = discover_rotation_workbook("Test Company", 2026, self.root)
        self.assertEqual(issues, [])
        self.assertEqual(discovered, preferred)
        self.assertNotEqual(discovered, root_path)
        context = build_document_context("test company", self.provider).context
        selected, issues = read_retiring_directors(preferred, 2025, context.directors)
        self.assertEqual(issues, [])
        self.assertEqual([person.name for person in selected], ["PERSON 1", "PERSON 2"])

    def test_missing_and_ambiguous_workbooks_block(self) -> None:
        path, issues = discover_rotation_workbook("Test Company", 2026, self.root)
        self.assertIsNone(path)
        self.assertEqual(issues[0].code, "rotation_workbook_missing")
        write_rotation(self.root, self.provider, ["PERSON 1"], filename="Retirement A.xlsx")
        write_rotation(self.root, self.provider, ["PERSON 1"], filename="Rotation B.xlsx")
        path, issues = discover_rotation_workbook("Test Company", 2026, self.root)
        self.assertIsNone(path)
        self.assertEqual(issues[0].code, "rotation_workbook_ambiguous")

    def test_missing_year_unmarked_and_mismatched_names_block(self) -> None:
        context = build_document_context("test company", self.provider).context
        missing_year = write_rotation(self.root, self.provider, ["PERSON 1"], year=2024)
        selected, issues = read_retiring_directors(missing_year, 2025, context.directors)
        self.assertEqual(selected, [])
        self.assertEqual(issues[0].code, "rotation_year_missing")

        unmarked = write_rotation(self.root, self.provider, [], filename="Unmarked Rotation.xlsx")
        selected, issues = read_retiring_directors(unmarked, 2025, context.directors)
        self.assertEqual(issues[0].code, "rotation_unmarked")

        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Director", 2025])
        sheet.append(["NOT A CURRENT DIRECTOR", "X"])
        mismatch = self.root / "Mismatch.xlsx"
        workbook.save(mismatch)
        workbook.close()
        selected, issues = read_retiring_directors(mismatch, 2025, context.directors)
        self.assertEqual(issues[0].code, "rotation_director_mismatch")


class Section90RenderingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_one_through_six_people_use_two_person_rows_without_markers(self) -> None:
        for count in range(1, 7):
            with self.subTest(count=count):
                provider = section90_provider(count)
                write_rotation(self.root, provider, [provider.company["Director1 Name"]])
                results = generate_documents(
                    ["test company"],
                    output_dir=self.root / f"output-{count}",
                    provider=provider,
                    clients_root=self.root,
                    section90_inputs={"test company": "THIRTEENTH"},
                )
                self.assertEqual(results[0].status, "generated", results[0].issues)
                output = results[0].output_path
                self.assertEqual(validate_output_docx(output, TEMPLATE_SECTION_90), [])
                document = Document(output)
                all_text = "\n".join(
                    [p.text for p in document.paragraphs]
                    + [c.text for t in document.tables for r in t.rows for c in r.cells]
                )
                self.assertNotIn("{{", all_text)
                self.assertNotIn("AUTHORITY TO FILE", all_text)
                self.assertNotIn("PAYMENT OF DIRECTORS", all_text)
                self.assertIn("THIRTEENTH ANNUAL GENERAL MEETING", all_text)
                people_grids = [table for table in document.tables if len(table.columns) == 3]
                self.assertEqual(len(people_grids), 3)
                self.assertTrue(
                    all(len(table.rows) == (count + 1) // 2 for table in people_grids)
                )
                attendance_tables = [
                    table
                    for table in document.tables
                    if len(table.columns) == 2
                    and table.cell(0, 0).text == "NAME"
                    and table.cell(0, 1).text == "SIGNATURE"
                ]
                self.assertEqual(len(attendance_tables), 1)
                attendance = attendance_tables[0]
                self.assertEqual(len(attendance.rows), 11)
                self.assertEqual(
                    [attendance.cell(row, 0).text for row in range(1, count + 1)],
                    [
                        provider.company[f"Member{index} Name"]
                        for index in range(1, count + 1)
                    ],
                )
                self.assertTrue(
                    all(not attendance.cell(row, 1).text for row in range(1, 11))
                )

    def test_prompt_callback_and_blank_generic_ordinal(self) -> None:
        provider = section90_provider(2)
        write_rotation(self.root, provider, ["PERSON 1"])
        calls: list[str] = []
        result = generate_documents(
            ["test company"],
            output_dir=self.root / "prompted",
            provider=provider,
            clients_root=self.root,
            ordinal_provider=lambda company: calls.append(company) or "THIRTEENTH",
        )[0]
        self.assertEqual(result.status, "generated", result.issues)
        self.assertEqual(calls, ["TEST COMPANY SDN. BHD."])

        generic = generate_documents(
            ["test company"],
            output_dir=self.root / "generic",
            provider=provider,
            clients_root=self.root,
            section90_inputs={"TEST COMPANY SDN. BHD.": ""},
        )[0]
        document = Document(generic.output_path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("NOTICE OF ANNUAL GENERAL MEETING", text)
        self.assertNotIn("THIRTEENTH", text)

    def test_article_5_without_to_completes_template_3_workflow(self) -> None:
        provider = section90_provider(2)
        provider.constitution[
            "DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"
        ] = SECTION_90_DWRS[0]
        write_rotation(self.root, provider, ["PERSON 1"])
        result = generate_documents(
            ["test company"],
            output_dir=self.root / "article-5-without-to",
            provider=provider,
            clients_root=self.root,
            section90_inputs={"test company": "THIRTEENTH"},
        )[0]
        self.assertEqual(result.status, "generated", result.issues)
        self.assertEqual(
            result.preview["selected_template_path"],
            str(SECTION_90_TEMPLATE),
        )
        self.assertEqual(
            validate_output_docx(result.output_path, TEMPLATE_SECTION_90),
            [],
        )

    def test_dry_run_reports_prompt_and_derived_inputs_without_pausing(self) -> None:
        provider = section90_provider(1)
        workbook = write_rotation(self.root, provider, ["PERSON 1"])
        result = generate_documents(
            ["test company"],
            dry_run=True,
            provider=provider,
            clients_root=self.root,
            ordinal_provider=lambda company: self.fail("dry-run must not prompt"),
        )[0]
        self.assertEqual(result.status, "valid", result.issues)
        self.assertTrue(result.preview["agm_ordinal_prompt_required"])
        self.assertEqual(result.preview["selected_template_path"], str(SECTION_90_TEMPLATE))
        self.assertEqual(result.preview["notice_date"], "2026-06-12")
        self.assertEqual(result.preview["rotation_workbook_path"], str(workbook))


if __name__ == "__main__":
    unittest.main()
