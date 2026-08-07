from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from document_generation.context_builder import build_document_context
from document_generation.generator import (
    DEFAULT_TEMPLATE,
    FIRST_AGM_TEMPLATE,
    SECTION_90_TEMPLATE,
    generate_documents,
    resolve_template_path,
)
from document_generation.models import Person
from document_generation.output_validation import validate_output_docx
from document_generation.rotation_reader import (
    discover_rotation_workbook,
    read_retiring_directors,
)
from document_generation.template_selection import (
    TEMPLATE_PARAGRAPH_15,
    TEMPLATE_SECTION_90,
    TEMPLATE_STANDARD,
    classify_dwr,
)

from document_generation.tests.test_generator import FakeProvider


PARAGRAPH_15 = (
    "DIRECTORS’ WRITTEN RESOLUTION PASSED PURSUANT TO PARAGRAPH 15 "
    "OF THE THIRD SCHEDULE OF THE COMPANIES ACT, 2016"
)
REGULATION_90 = (
    "DIRECTORS� WRITTEN RESOLUTION PASSED PURSUANT TO REGULATION 90 "
    "(TABLE A) OF THE COMPANY�S ARTICLES OF ASSOCIATION (�THE CONSTITUTION�)"
)


def section90_provider(count: int = 3) -> FakeProvider:
    provider = FakeProvider(
        member_count=count,
        fee="not-a-number",
        financial_year_start="2025-04-03",
    )
    provider.constitution["DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"] = REGULATION_90
    provider.constitution["MEMBER WRITTEN RESOLUTION (MWR Statutory)"] = ""
    return provider


def write_rotation(
    clients_root: Path,
    provider: FakeProvider,
    checked_names: list[str],
    *,
    year: int = 2025,
    circulation_subfolder: bool = True,
    filename: str = "Retirement Rotation.xlsx",
) -> Path:
    base = clients_root / provider.company["Folder"] / "AGM"
    if circulation_subfolder:
        base = base / "2026"
    base.mkdir(parents=True, exist_ok=True)
    path = base / filename
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["No.", "Director", year])
    for index in range(1, provider.member_count + 1):
        name = provider.company[f"Director{index} Name"]
        sheet.append([index, name, "✓" if name in checked_names else ""])
    workbook.save(path)
    workbook.close()
    return path


class DwrClassifierTests(unittest.TestCase):
    def test_classifies_paragraph_15_and_regulation_90_variants(self) -> None:
        self.assertEqual(classify_dwr(PARAGRAPH_15), TEMPLATE_PARAGRAPH_15)
        variants = (
            REGULATION_90,
            "Directors' Written Resolution pursuant to Regulation 90, Table A, "
            "of the Company's Articles of Association (Companies Constitution)",
            "DIRECTORS’ WRITTEN RESOLUTION REGULATION 90 (TABLE A) "
            "ARTICLES OF ASSOCIATION",
        )
        for value in variants:
            with self.subTest(value=value):
                self.assertEqual(classify_dwr(value), TEMPLATE_SECTION_90)

    def test_unknown_dwr_falls_back_to_template_1_even_for_irregular_period(self) -> None:
        provider = FakeProvider(financial_year_start="2025-04-03")
        provider.constitution["DIRECTOR WRITTEN RESOLUTION (DWR Statutory)"] = "UNKNOWN CLAUSE"
        result = build_document_context("test company", provider)
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(result.context.template_family, TEMPLATE_STANDARD)
        self.assertEqual(resolve_template_path(result.context), DEFAULT_TEMPLATE)

    def test_paragraph_15_uses_irregular_period_rule(self) -> None:
        result = build_document_context(
            "test company",
            FakeProvider(financial_year_start="2025-04-03"),
        )
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(resolve_template_path(result.context), FIRST_AGM_TEMPLATE)

    def test_regulation_90_ignores_period_length_and_blank_mwr(self) -> None:
        result = build_document_context("test company", section90_provider())
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(result.context.template_family, TEMPLATE_SECTION_90)
        self.assertEqual(resolve_template_path(result.context), SECTION_90_TEMPLATE)
        self.assertEqual(result.context.fee_allocations, [])


class RotationWorkbookTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.provider = section90_provider(3)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_prefers_circulation_year_subfolder_and_reads_multiple_names(self) -> None:
        root_path = write_rotation(
            self.root,
            self.provider,
            ["PERSON 3"],
            circulation_subfolder=False,
        )
        preferred = write_rotation(
            self.root,
            self.provider,
            ["PERSON 1", "PERSON 2"],
        )
        discovered, issues = discover_rotation_workbook("Test Company", 2026, self.root)
        self.assertEqual(issues, [])
        self.assertEqual(discovered, preferred)
        self.assertNotEqual(discovered, root_path)
        context = build_document_context("test company", self.provider).context
        selected, issues = read_retiring_directors(preferred, 2025, context.directors)
        self.assertEqual(issues, [])
        self.assertEqual([person.name for person in selected], ["PERSON 1", "PERSON 2"])

    def test_missing_and_ambiguous_workbooks_block(self) -> None:
        path, issues = discover_rotation_workbook("Test Company", 2026, self.root)
        self.assertIsNone(path)
        self.assertEqual(issues[0].code, "rotation_workbook_missing")
        write_rotation(self.root, self.provider, ["PERSON 1"], filename="Retirement A.xlsx")
        write_rotation(self.root, self.provider, ["PERSON 1"], filename="Rotation B.xlsx")
        path, issues = discover_rotation_workbook("Test Company", 2026, self.root)
        self.assertIsNone(path)
        self.assertEqual(issues[0].code, "rotation_workbook_ambiguous")

    def test_missing_year_unmarked_and_mismatched_names_block(self) -> None:
        context = build_document_context("test company", self.provider).context
        missing_year = write_rotation(self.root, self.provider, ["PERSON 1"], year=2024)
        selected, issues = read_retiring_directors(missing_year, 2025, context.directors)
        self.assertEqual(selected, [])
        self.assertEqual(issues[0].code, "rotation_year_missing")

        unmarked = write_rotation(self.root, self.provider, [], filename="Unmarked Rotation.xlsx")
        selected, issues = read_retiring_directors(unmarked, 2025, context.directors)
        self.assertEqual(issues[0].code, "rotation_unmarked")

        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Director", 2025])
        sheet.append(["NOT A CURRENT DIRECTOR", "X"])
        mismatch = self.root / "Mismatch.xlsx"
        workbook.save(mismatch)
        workbook.close()
        selected, issues = read_retiring_directors(mismatch, 2025, context.directors)
        self.assertEqual(issues[0].code, "rotation_director_mismatch")


class Section90RenderingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_one_through_six_people_use_two_person_rows_without_markers(self) -> None:
        for count in range(1, 7):
            with self.subTest(count=count):
                provider = section90_provider(count)
                write_rotation(self.root, provider, [provider.company["Director1 Name"]])
                results = generate_documents(
                    ["test company"],
                    output_dir=self.root / f"output-{count}",
                    provider=provider,
                    clients_root=self.root,
                    section90_inputs={"test company": "THIRTEENTH"},
                )
                self.assertEqual(results[0].status, "generated", results[0].issues)
                output = results[0].output_path
                self.assertEqual(validate_output_docx(output, TEMPLATE_SECTION_90), [])
                document = Document(output)
                all_text = "\n".join(
                    [p.text for p in document.paragraphs]
                    + [c.text for t in document.tables for r in t.rows for c in r.cells]
                )
                self.assertNotIn("{{", all_text)
                self.assertNotIn("AUTHORITY TO FILE", all_text)
                self.assertNotIn("PAYMENT OF DIRECTORS", all_text)
                self.assertIn("THIRTEENTH ANNUAL GENERAL MEETING", all_text)
                people_grids = [table for table in document.tables if len(table.columns) == 3]
                self.assertEqual(len(people_grids), 3)
                self.assertTrue(
                    all(len(table.rows) == (count + 1) // 2 for table in people_grids)
                )
                attendance_tables = [
                    table
                    for table in document.tables
                    if len(table.columns) == 2
                    and table.cell(0, 0).text == "NAME"
                    and table.cell(0, 1).text == "SIGNATURE"
                ]
                self.assertEqual(len(attendance_tables), 1)
                attendance = attendance_tables[0]
                self.assertEqual(len(attendance.rows), 11)
                self.assertEqual(
                    [attendance.cell(row, 0).text for row in range(1, count + 1)],
                    [
                        provider.company[f"Member{index} Name"]
                        for index in range(1, count + 1)
                    ],
                )
                self.assertTrue(
                    all(not attendance.cell(row, 1).text for row in range(1, 11))
                )

    def test_prompt_callback_and_blank_generic_ordinal(self) -> None:
        provider = section90_provider(2)
        write_rotation(self.root, provider, ["PERSON 1"])
        calls: list[str] = []
        result = generate_documents(
            ["test company"],
            output_dir=self.root / "prompted",
            provider=provider,
            clients_root=self.root,
            ordinal_provider=lambda company: calls.append(company) or "THIRTEENTH",
        )[0]
        self.assertEqual(result.status, "generated", result.issues)
        self.assertEqual(calls, ["TEST COMPANY SDN. BHD."])

        generic = generate_documents(
            ["test company"],
            output_dir=self.root / "generic",
            provider=provider,
            clients_root=self.root,
            section90_inputs={"TEST COMPANY SDN. BHD.": ""},
        )[0]
        document = Document(generic.output_path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("NOTICE OF ANNUAL GENERAL MEETING", text)
        self.assertNotIn("THIRTEENTH", text)

    def test_dry_run_reports_prompt_and_derived_inputs_without_pausing(self) -> None:
        provider = section90_provider(1)
        workbook = write_rotation(self.root, provider, ["PERSON 1"])
        result = generate_documents(
            ["test company"],
            dry_run=True,
            provider=provider,
            clients_root=self.root,
            ordinal_provider=lambda company: self.fail("dry-run must not prompt"),
        )[0]
        self.assertEqual(result.status, "valid", result.issues)
        self.assertTrue(result.preview["agm_ordinal_prompt_required"])
        self.assertEqual(result.preview["selected_template_path"], str(SECTION_90_TEMPLATE))
        self.assertEqual(result.preview["notice_date"], "2026-06-12")
        self.assertEqual(result.preview["rotation_workbook_path"], str(workbook))


if __name__ == "__main__":
    unittest.main()
