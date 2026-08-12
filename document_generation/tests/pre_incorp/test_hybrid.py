from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from document_generation.pre_incorp_generation.cli import main
from document_generation.pre_incorp_generation.generator import generate_pre_incorp_documents
from document_generation.pre_incorp_generation.interactive import (
    _edit_roster,
    _resolve_field,
    complete_interactively,
)
from document_generation.pre_incorp_generation.models import DraftField, SourceCandidate
from document_generation.pre_incorp_generation.preparation import (
    prepare_pre_incorp_generation,
)
from document_generation.pre_incorp_generation.renderer import format_identification


class HybridProvider:
    def get_company(self, company_name: str) -> list[dict]:
        return [
            {
                "Company Name": "HOSAY 3 BAKERY SDN. BHD.",
                "Reg No": "202401051970 (1597813-X)",
                "Incorporate Date": "09/12/2024",
                "Director1 Name": "TAM WEE SEONG",
                "Director1 IC": "811229025695",
                "Director1 ID Type": "MYKAD",
                "Director1 DOB": "1981-12-29",
                "Director1 Nationality": "MALAYSIA",
                "Director1 Race": "CHINESE",
                "Director1 Residential Address": "CURRENT RESIDENTIAL ADDRESS",
                "Director1 Service Address": "CURRENT SERVICE ADDRESS",
                "Director1 Business Occupation": "DIRECTOR",
                "Member1 Name": "TAM WEE SEONG",
                "Member1 ID No": "811229025695",
                "Member1 Shares": "2",
            }
        ]

    def get_statutory_documents(self, company_name: str) -> list[dict]:
        payload = {
            "directors": [
                {
                    "Name": "TAM WEE SEONG",
                    "IC": "811229025695",
                    "DOB": "29/12/1981",
                    "Nationality": "MALAYSIA",
                    "Race": "CHINESE",
                    "Residential": "INCORPORATION RESIDENTIAL TAM",
                    "Service Address": "tam@example.test",
                },
                {
                    "Name": "TAN HUI KEE",
                    "IC": "930412025193",
                    "DOB": "12/04/1993",
                    "Nationality": "MALAYSIA",
                    "Race": "CHINESE",
                    "Residential": "INCORPORATION RESIDENTIAL TAN",
                    "Service Address": "tan@example.test",
                },
            ],
            "members": [
                {
                    "Name": "TAM WEE SEONG",
                    "ID No": "811229025695",
                    "ID Type": "NRIC",
                    "Shares": 1,
                },
                {
                    "Name": "TAN HUI KEE",
                    "ID No": "930412025193",
                    "ID Type": "NRIC",
                    "Shares": 1,
                },
            ],
        }
        return [
            {
                "Section": "S14",
                "Status": "VALID",
                "EffectiveDate": "2024-12-09",
                "ParsedJSON": json.dumps(payload),
            }
        ]

    def get_ebos_events(self, company_name: str) -> list[dict]:
        common = {
            "Company Name": "HOSAY 3 BAKERY SDN. BHD.",
            "Company No": "202401051970 (1597813-X)",
            "Date of Becoming BO": "09/12/2024",
            "Nationality": "MALAYSIA",
            "Race": "CHINESE",
            "Designation": "DIRECTOR",
        }
        return [
            {
                **common,
                "Name": "TAM WEE SEONG",
                "IC": "811229025695",
                "DOB": "29/12/1981",
                "Residential Address": "INCORPORATION RESIDENTIAL TAM",
                "Email": "tam@example.test",
                "Contact No": "0175717592",
                "Date Received": "18/12/2024",
            },
            {
                **common,
                "Name": "TAM WEE SEONG",
                "IC": "811229025695",
                "DOB": "29/12/1981",
                "Residential Address": "CURRENT RESIDENTIAL ADDRESS",
                "Email": "tam@example.test",
                "Contact No": "0126220099",
                "Date Received": "25/05/2025",
            },
            {
                **common,
                "Name": "TAN HUI KEE",
                "IC": "930412025193",
                "DOB": "12/04/1993",
                "Residential Address": "INCORPORATION RESIDENTIAL TAN",
                "Email": "tan@example.test",
                "Contact No": "0175962284",
                "Date Received": "18/12/2024",
            },
        ]


class CurrentOnlyProvider(HybridProvider):
    def get_statutory_documents(self, company_name: str) -> list[dict]:
        return []

    def get_ebos_events(self, company_name: str) -> list[dict]:
        return []


def document_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(values)


class PreparationTests(unittest.TestCase):
    def test_notice_identification_type_codes(self) -> None:
        expected = {
            "MYKAD": "B",
            "NRIC": "B",
            "PASSPORT": "P",
            "RED IC": "R",
            "MILITARY ID": "Z",
            "POLICE ID": "M",
        }
        for id_type, code in expected.items():
            with self.subTest(id_type=id_type):
                self.assertEqual(format_identification("123", id_type), f"123 ({code})")

    def test_detects_initial_directors_shares_and_incorporation_phone(self) -> None:
        result = prepare_pre_incorp_generation("hosay 3", HybridProvider())
        self.assertTrue(result.found, result.issues)
        draft = result.draft
        self.assertEqual([person.name.value for person in draft.directors], ["TAM WEE SEONG", "TAN HUI KEE"])
        self.assertEqual([str(person.shares.value) for person in draft.directors], ["1", "1"])
        self.assertEqual(draft.directors[0].phone.value, "0175717592")
        self.assertEqual(draft.directors[0].phone.status, "conflicting")
        self.assertEqual(
            [candidate.value for candidate in draft.directors[0].phone.candidates],
            ["0175717592", "0126220099"],
        )
        self.assertNotIn("date_of_birth", draft.directors[0].required_inputs())

    def test_current_director_fallback_is_provisional(self) -> None:
        draft = prepare_pre_incorp_generation("hosay 3", CurrentOnlyProvider()).draft
        self.assertEqual(len(draft.directors), 1)
        self.assertEqual(draft.directors[0].roster_status, "provisional")
        self.assertIn("director_roster_confirmation", draft.directors[0].required_inputs())

    def test_generation_uses_completed_draft_and_notice_id_code(self) -> None:
        draft = prepare_pre_incorp_generation("hosay 3", HybridProvider()).draft
        draft.directors[0].service_address.set_user_value("SERVICE ADDRESS TAM")
        draft.directors[0].phone.set_user_value("0175717592")
        draft.directors[1].service_address.set_user_value("SERVICE ADDRESS TAN")
        draft.directors[1].occupation.set_user_value("DIRECTOR")
        with tempfile.TemporaryDirectory() as directory:
            result = generate_pre_incorp_documents(
                "hosay 3",
                reference_no="REF-HYBRID",
                output_dir=Path(directory),
                draft=draft,
                confirmed=True,
            )
            self.assertEqual(result.status, "generated", result.issues)
            self.assertEqual(len(result.output_paths), 4)
            tam_notice = next(
                path
                for path in result.output_paths
                if path.name.startswith("TAM WEE SEONG") and "Notice" in path.name
            )
            text = document_text(tam_notice)
            self.assertIn("811229025695 (B)", text)
            self.assertIn("1", text)

    def test_dry_run_reports_required_inputs_without_writing(self) -> None:
        draft = prepare_pre_incorp_generation("hosay 3", HybridProvider()).draft
        with tempfile.TemporaryDirectory() as directory:
            result = generate_pre_incorp_documents(
                "hosay 3",
                output_dir=Path(directory),
                dry_run=True,
                draft=draft,
            )
            self.assertEqual(result.status, "input_required")
            self.assertIn("director[1].phone", result.preview["draft"]["required_inputs"])
            self.assertEqual(list(Path(directory).rglob("*.docx")), [])


class InteractiveTests(unittest.TestCase):
    def test_cancel_stops_before_generation(self) -> None:
        draft = prepare_pre_incorp_generation("hosay 3", HybridProvider()).draft
        status, reference = complete_interactively(draft, None, lambda _: "c", lambda _: None)
        self.assertEqual((status, reference), ("canceled", ""))

    def test_final_rejection_cancels_completed_draft(self) -> None:
        draft = prepare_pre_incorp_generation("hosay 3", HybridProvider()).draft
        draft.directors[0].residential_address.status = "detected"
        draft.directors[0].service_address.set_user_value("SERVICE TAM")
        draft.directors[0].phone.status = "detected"
        draft.directors[1].service_address.set_user_value("SERVICE TAN")
        draft.directors[1].occupation.set_user_value("DIRECTOR")
        answers = iter(["a", "n"])
        status, reference = complete_interactively(
            draft,
            "REF",
            lambda _: next(answers),
            lambda _: None,
        )
        self.assertEqual((status, reference), ("canceled", "REF"))

    def test_roster_editor_adds_removes_and_corrects_directors(self) -> None:
        draft = prepare_pre_incorp_generation("hosay 3", HybridProvider()).draft
        answers = iter(
            [
                "remove 2",
                "add",
                "NEW DIRECTOR",
                "900101025555",
                "edit 1",
                "TAM WEE SEONG CORRECTED",
                *([""] * 11),
                "done",
            ]
        )
        _edit_roster(draft, lambda _: next(answers), lambda _: None)
        self.assertEqual(len(draft.directors), 2)
        self.assertEqual(draft.directors[0].name.value, "TAM WEE SEONG CORRECTED")
        self.assertEqual(draft.directors[1].name.value, "NEW DIRECTOR")
        self.assertEqual(draft.directors[1].id_number.value, "900101025555")

    def test_conflict_choices_and_invalid_share_retry(self) -> None:
        phone = DraftField(
            value="0171111111",
            source="EBOS",
            status="conflicting",
            candidates=[
                SourceCandidate("0171111111", "EBOS"),
                SourceCandidate("0122222222", "EBOS"),
            ],
        )
        phone_answers = iter(["invalid", "2"])
        _resolve_field("phone", phone, lambda _: next(phone_answers), lambda _: None)
        self.assertEqual(phone.value, "0122222222")
        self.assertEqual(phone.status, "user-supplied")

        shares = DraftField()
        share_answers = iter(["not-a-number", "-1", "25"])
        _resolve_field("shares", shares, lambda _: next(share_answers), lambda _: None)
        self.assertEqual(str(shares.value), "25")

    def test_non_tty_cli_reports_input_required_without_prompting(self) -> None:
        prepared = prepare_pre_incorp_generation("hosay 3", HybridProvider())
        with patch(
            "document_generation.pre_incorp_generation.cli.prepare_pre_incorp_generation",
            return_value=prepared,
        ), patch(
            "document_generation.pre_incorp_generation.cli.sys.stdin.isatty",
            return_value=False,
        ), patch(
            "document_generation.pre_incorp_generation.cli.print"
        ), patch("document_generation.pre_incorp_generation.cli.input") as prompt:
            exit_code = main(["--company", "hosay 3", "--reference-no", "REF"])
        self.assertEqual(exit_code, 1)
        prompt.assert_not_called()


if __name__ == "__main__":
    unittest.main()
