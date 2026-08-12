from __future__ import annotations

import hashlib
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from document_generation.pre_incorp_generation.cli import main
from document_generation.pre_incorp_generation.context_builder import (
    build_pre_incorp_context,
)
from document_generation.pre_incorp_generation.generator import (
    NOTICE_TEMPLATE,
    S201_TEMPLATE,
    generate_pre_incorp_documents,
)
from document_generation.pre_incorp_generation.preparation import (
    prepare_pre_incorp_generation,
)


class FakeProvider:
    def __init__(self, director_count: int = 2) -> None:
        row = {
            "Company Name": "EXAMPLE FOODS SDN. BHD.",
            "Reg No": "202501000001 (1600001-A)",
            "Incorporate Date": "10/01/2025",
        }
        for index in range(1, director_count + 1):
            name = f"EXAMPLE DIRECTOR {index}"
            row.update(
                {
                    f"Director{index} Name": name,
                    f"Director{index} IC": f"80010102000{index}",
                    f"Director{index} ID Type": "MYKAD",
                    f"Director{index} DOB": f"198{index}-01-0{index}",
                    f"Director{index} Nationality": "MALAYSIA",
                    f"Director{index} Race": "CHINESE",
                    f"Director{index} Residential Address": (
                        f"NO {index}, A DELIBERATELY LONG RESIDENTIAL ADDRESS, "
                        "TAMAN EXAMPLE, 08000 SUNGAI PETANI KEDAH MALAYSIA"
                    ),
                    f"Director{index} Service Address": (
                        f"SUITE {index}, BUSINESS CENTRE, 08000 SUNGAI PETANI KEDAH MALAYSIA"
                    ),
                    f"Director{index} Business Occupation": "DIRECTOR",
                    f"Director{index} Email": f"director{index}@example.test",
                    f"Director{index} Contact No": f"012000000{index}",
                    f"Member{index} Name": name,
                    f"Member{index} Shares": str(index * 10),
                }
            )
        self.row = row

    def get_company(self, company_name: str) -> list[dict]:
        return [self.row]


def document_text(path: Path) -> str:
    document = Document(str(path))
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(values)


class ContextTests(unittest.TestCase):
    def test_builds_all_directors_and_derives_declaration_date(self) -> None:
        result = build_pre_incorp_context("example foods", FakeProvider(2))
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(result.context.declaration_date.isoformat(), "2025-01-09")
        self.assertEqual([person.shares for person in result.context.directors], [10, 20])

    def test_any_missing_required_director_field_blocks_company(self) -> None:
        provider = FakeProvider(2)
        provider.row["Director2 Email"] = ""
        with tempfile.TemporaryDirectory() as directory:
            result = generate_pre_incorp_documents(
                "example foods",
                reference_no="REF-1",
                output_dir=Path(directory),
                provider=provider,
                confirmed=True,
            )
            self.assertEqual(result.status, "invalid")
            self.assertEqual(list(Path(directory).rglob("*.docx")), [])
            self.assertIn("missing_director_data", {issue.code for issue in result.issues})


class GenerationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.output = Path(self.temporary.name)

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def test_generates_two_separate_documents_per_director(self) -> None:
        result = generate_pre_incorp_documents(
            "example foods",
            reference_no="2025B000001",
            output_dir=self.output,
            provider=FakeProvider(2),
            confirmed=True,
        )
        self.assertEqual(result.status, "generated", result.issues)
        self.assertEqual(len(result.output_paths), 4)
        self.assertTrue(all(path.is_file() for path in result.output_paths))
        self.assertTrue(all(path.parent.name == "EXAMPLE FOODS SDN. BHD" for path in result.output_paths))

        for path in result.output_paths:
            text = document_text(path)
            director_number = 1 if path.name.startswith("EXAMPLE DIRECTOR 1") else 2
            self.assertIn(f"EXAMPLE DIRECTOR {director_number}", text)
            self.assertIn("EXAMPLE FOODS SDN. BHD.", text)
            self.assertNotIn("TAN HUI KEE", text)
            self.assertNotIn("930412-02-5193", text)
            if "S201" in path.name:
                self.assertIn("2025B000001", text)
                self.assertIn("9 January 2025", text)
                self.assertIn(f"012000000{director_number}", text)
                self.assertNotIn("DIRECTOR’S NOTICE", text)
            else:
                self.assertIn(f"{director_number * 10}", text)
                self.assertIn("MALAYSIAN/CHINESE", text)
                self.assertIn(f"80010102000{director_number} (B)", text)
                self.assertNotIn("9 January 2025", text)
                self.assertNotIn("Section 201", text)
                document = Document(str(path))
                detail_rows = document.tables[0].rows
                for row_index in (7, 11, 13, 15, 17, 19, 21, 23):
                    self.assertFalse(detail_rows[row_index].cells[-1].text.startswith(":"))

    def test_programmatic_generation_requires_confirmation_for_provisional_values(self) -> None:
        result = generate_pre_incorp_documents(
            "example foods",
            reference_no="REF-1",
            output_dir=self.output,
            provider=FakeProvider(1),
        )
        self.assertEqual(result.status, "input_required")
        self.assertEqual({issue.code for issue in result.issues}, {"draft_confirmation_required"})
        self.assertEqual(list(self.output.rglob("*.docx")), [])

    def test_collision_blocks_complete_batch_and_overwrite_replaces(self) -> None:
        provider = FakeProvider(1)
        first = generate_pre_incorp_documents(
            "example foods",
            "REF-1",
            self.output,
            provider=provider,
            confirmed=True,
        )
        self.assertEqual(first.status, "generated", first.issues)
        blocked = generate_pre_incorp_documents(
            "example foods",
            "REF-2",
            self.output,
            provider=provider,
            confirmed=True,
        )
        self.assertEqual(blocked.status, "invalid")
        self.assertEqual({issue.code for issue in blocked.issues}, {"output_exists"})
        overwritten = generate_pre_incorp_documents(
            "example foods",
            "REF-2",
            self.output,
            overwrite=True,
            provider=provider,
            confirmed=True,
        )
        self.assertEqual(overwritten.status, "generated", overwritten.issues)
        s201 = next(path for path in overwritten.output_paths if "S201" in path.name)
        self.assertIn("REF-2", document_text(s201))

    def test_dry_run_does_not_require_reference_or_create_files(self) -> None:
        result = generate_pre_incorp_documents(
            "example foods",
            output_dir=self.output,
            dry_run=True,
            provider=FakeProvider(1),
        )
        self.assertEqual(result.status, "input_required", result.issues)
        self.assertTrue(result.preview["reference_no_prompt_required"])
        self.assertIn("draft", result.preview)
        self.assertEqual(list(self.output.rglob("*.docx")), [])


class CliTests(unittest.TestCase):
    def test_cli_retrieves_before_reference_and_never_asks_director_count(self) -> None:
        prepared = prepare_pre_incorp_generation("EXAMPLE FOODS", FakeProvider(1))
        prompts: list[str] = []

        def answer(prompt: str) -> str:
            prompts.append(prompt)
            answers = {
                "Director list": "a",
                "Reference No.": "REF-PROMPTED",
                "Press Enter": "",
                "Generate two": "y",
            }
            return next(value for key, value in answers.items() if key in prompt)

        with patch(
            "document_generation.pre_incorp_generation.cli.input",
            side_effect=answer,
        ) as prompt, patch(
            "document_generation.pre_incorp_generation.cli.print"
        ), patch(
            "document_generation.pre_incorp_generation.cli.sys.stdin.isatty",
            return_value=True,
        ), patch(
            "document_generation.pre_incorp_generation.cli.prepare_pre_incorp_generation",
            return_value=prepared,
        ), patch(
            "document_generation.pre_incorp_generation.cli.generate_pre_incorp_documents"
        ) as generate:
            generate.return_value.status = "generated"
            generate.return_value.to_dict.return_value = {"status": "generated"}
            exit_code = main(["--company", "EXAMPLE FOODS", "--output-dir", str(Path("output"))])
        self.assertEqual(exit_code, 0)
        self.assertGreaterEqual(prompt.call_count, 4)
        self.assertIn("Director list", prompts[0])
        self.assertFalse(any("number of director" in value.lower() for value in prompts))
        self.assertEqual(generate.call_args.kwargs["reference_no"], "REF-PROMPTED")
        self.assertIs(generate.call_args.kwargs["draft"], prepared.draft)

    def test_cli_dry_run_never_prompts(self) -> None:
        prepared = prepare_pre_incorp_generation("EXAMPLE FOODS", FakeProvider(1))
        with patch(
            "document_generation.pre_incorp_generation.cli.input"
        ) as prompt, patch(
            "document_generation.pre_incorp_generation.cli.print"
        ), patch(
            "document_generation.pre_incorp_generation.cli.prepare_pre_incorp_generation",
            return_value=prepared,
        ), patch(
            "document_generation.pre_incorp_generation.cli.generate_pre_incorp_documents"
        ) as generate:
            generate.return_value.status = "input_required"
            generate.return_value.to_dict.return_value = {"status": "input_required"}
            exit_code = main(["--company", "EXAMPLE FOODS", "--dry-run"])
        self.assertEqual(exit_code, 1)
        prompt.assert_not_called()
        self.assertTrue(generate.call_args.kwargs["dry_run"])


class TemplateContractTests(unittest.TestCase):
    def test_retained_template_hashes(self) -> None:
        expected = {
            S201_TEMPLATE: "98ffb593cc0341e85fdf94c06a3175aee92c7123a639f7ba86d4ad7d1d60989e",
            NOTICE_TEMPLATE: "4f855fe3b840c8f4298a2c0a7fbe3c674bb0977f9028c4a787f80c24fd4efcf2",
        }
        for path, digest in expected.items():
            with self.subTest(path=path):
                self.assertEqual(hashlib.sha256(path.read_bytes()).hexdigest(), digest)

    def test_template_structures_are_not_interchanged(self) -> None:
        s201 = Document(str(S201_TEMPLATE))
        notice = Document(str(NOTICE_TEMPLATE))
        self.assertEqual((len(s201.sections), len(s201.tables)), (1, 3))
        self.assertEqual((len(notice.sections), len(notice.tables)), (2, 3))
        self.assertNotEqual(s201.sections[0].orientation, notice.sections[1].orientation)


if __name__ == "__main__":
    unittest.main()
