from __future__ import annotations

import hashlib
import tempfile
import unittest
from copy import deepcopy
from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document

from csai_langchain.config.settings import NEW_INCORP_DB
from document_generation.new_incorp_generation import (
    generate_new_incorp_documents,
    prepare_new_incorp_generation,
)
from document_generation.new_incorp_generation.audit_templates import audit_templates
from document_generation.new_incorp_generation.generator import TEMPLATE_DIR
from document_generation.new_incorp_generation.renderer import render_template


class FakeProvider:
    def __init__(self, rows):
        self.rows = rows

    def get_company(self, company_name: str):
        return deepcopy(self.rows)


def _person(row: dict, prefix: str, index: int, name: str, identity: str, shares: str | None = None) -> None:
    stem = f"{prefix}{index} "
    row.update(
        {
            stem + "Name": name,
            stem + "ID Type": "NRIC",
            stem + "Identification No": identity,
            stem + "Nationality": "MALAYSIAN",
            stem + "Race": "CHINESE",
            stem + "Gender": "MALE" if index % 2 else "FEMALE",
            stem + "DOB": f"198{index}-01-02",
            stem + "Address": f"{index} SAMPLE ROAD 08000 SUNGAI PETANI KEDAH MALAYSIA",
            stem + "Email": f"person{index}@example.test",
            stem + "Contact No": f"01200000{index:02d}",
            stem + "Business Occupation": "DIRECTOR",
        }
    )
    if shares is not None:
        row[stem + "Number of Shares"] = shares
        row[stem + "Class of Share"] = "Ordinary"
        row[stem + "Price per Share"] = "1.00"


def _current_director(row: dict, index: int, name: str, identity: str) -> None:
    stem = f"Director{index} "
    row.update(
        {
            stem + "Name": name,
            stem + "IC": identity,
            stem + "ID Type": "MYKAD",
            stem + "Nationality": "MALAYSIA",
            stem + "Race": "CHINESE",
            stem + "Gender": "MALE" if index % 2 else "FEMALE",
            stem + "DOB": f"198{index}-01-02",
            stem + "Residential Address": f"{index} SAMPLE ROAD 08000 SUNGAI PETANI KEDAH MALAYSIA",
            stem + "Email": f"person{index}@example.test",
            stem + "Contact No": f"01200000{index:02d}",
            stem + "Business Occupation": "DIRECTOR",
        }
    )


def sample_row(directors: int = 2, members: int = 2) -> dict:
    row = {
        "Company Name": "EXAMPLE VENTURES SDN. BHD.",
        "Reg No": "202601000001 (1700001-A)",
        "S14 Proposed Name": "EXAMPLE VENTURES SDN. BHD.",
        "S14 Registration No": "202601000001 (1700001-A)",
        "S14 Incorporation Date": "2026-01-15",
        "S14 Registered Address": "10 REGISTERED ROAD 08000 SUNGAI PETANI KEDAH MALAYSIA",
        "S14 Business Address": "20 BUSINESS ROAD 08000 SUNGAI PETANI KEDAH MALAYSIA",
    }
    people = max(directors, members)
    for index in range(1, people + 1):
        name = f"PERSON {index}"
        identity = f"8001010200{index:02d}"
        if index <= directors:
            _person(row, "S14 Director", index, name, identity)
            _current_director(row, index, name, identity)
        if index <= members:
            _person(row, "S14 Member", index, name, identity, "1")
    return row


class ContextTests(unittest.TestCase):
    def test_not_found_and_ambiguous_lookup_block(self):
        missing = prepare_new_incorp_generation("missing", provider=FakeProvider([]))
        ambiguous = prepare_new_incorp_generation("ambiguous", provider=FakeProvider([sample_row(), sample_row()]))
        self.assertEqual(missing.issues[0].code, "source_not_found")
        self.assertEqual(ambiguous.issues[0].code, "source_ambiguous")

    def test_current_director_roster_excludes_historical_s14_directors(self):
        row = sample_row()
        row["S14 Director3 Name"] = "HISTORICAL DIRECTOR"
        row["S14 Director3 ID Type"] = "NRIC"
        row["S14 Director3 Identification No"] = "900101020099"
        row["Director1 Email"] = "current-director@example.test"
        result = prepare_new_incorp_generation(row["Company Name"], provider=FakeProvider([row]))
        self.assertTrue(result.valid, result.issues)
        self.assertEqual([person.name for person in result.context.directors], ["PERSON 1", "PERSON 2"])
        self.assertEqual(result.context.directors[0].email, "current-director@example.test")
        self.assertEqual(result.context.total_subscriber_shares, Decimal("2"))

    def test_corporate_member_and_missing_required_data_block_batch(self):
        corporate = sample_row()
        corporate["S14 Member1 ID Type"] = "COMPANY REGISTRATION"
        result = prepare_new_incorp_generation("x", provider=FakeProvider([corporate]))
        self.assertIn("unsupported_corporate_member", {issue.code for issue in result.issues})

        missing = sample_row()
        missing["S14 Director1 Email"] = ""
        missing["S14 Member1 Email"] = ""
        missing["Director1 Email"] = ""
        result = prepare_new_incorp_generation("x", provider=FakeProvider([missing]))
        self.assertIn("missing_person_data", {issue.code for issue in result.issues})

    def test_bo_percentage_threshold_and_historical_record(self):
        row = sample_row(directors=1, members=5)
        row["BO1 Name"] = "PERSON 1"
        row["BO1 Identification No"] = row["S14 Member1 Identification No"]
        row["BO1 Date of Becoming BO"] = "2025-01-01"
        row["BO1 Date of Cessation"] = "2026-12-31"
        result = prepare_new_incorp_generation("x", provider=FakeProvider([row]))
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(result.context.members[0].direct_percentage, Decimal("20"))
        self.assertTrue(result.context.members[0].beneficial_owner)
        self.assertEqual(result.context.members[0].becoming_bo_date, date(2025, 1, 1))

    def test_repository_is_read_only_for_real_lookup(self):
        before = hashlib.sha256(NEW_INCORP_DB.read_bytes()).hexdigest()
        result = prepare_new_incorp_generation("HOSAY 3 BAKERY SDN. BHD.")
        after = hashlib.sha256(NEW_INCORP_DB.read_bytes()).hexdigest()
        self.assertTrue(result.valid, result.issues)
        self.assertEqual(before, after)
        self.assertEqual([person.name for person in result.context.directors], ["TAM WEE SEONG"])
        self.assertEqual(result.context.document_count, 13)


class GenerationTests(unittest.TestCase):
    def test_template_fidelity_contract(self):
        result = audit_templates()
        self.assertTrue(result["valid"], result["errors"])

    def test_dry_run_count_generation_and_collision(self):
        provider = FakeProvider([sample_row(directors=3, members=6)])
        with tempfile.TemporaryDirectory() as directory:
            output_root = Path(directory)
            dry = generate_new_incorp_documents("EXAMPLE", output_root, dry_run=True, provider=provider)
            self.assertEqual(dry.status, "dry_run")
            self.assertEqual(len(dry.output_paths), 23)
            self.assertFalse((output_root / "EXAMPLE VENTURES SDN. BHD").exists())

            generated = generate_new_incorp_documents("EXAMPLE", output_root, provider=provider)
            self.assertEqual(generated.status, "generated", generated.issues)
            self.assertEqual(len(list((output_root / "EXAMPLE VENTURES SDN. BHD").glob("*.docx"))), 23)
            collision = generate_new_incorp_documents("EXAMPLE", output_root, provider=provider)
            self.assertEqual(collision.status, "collision")

    def test_odd_director_signature_grid_keeps_blank_right_cell(self):
        prepared = prepare_new_incorp_generation("EXAMPLE", provider=FakeProvider([sample_row(directors=3, members=1)]))
        self.assertTrue(prepared.valid, prepared.issues)
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "policy.docx"
            render_template(TEMPLATE_DIR / "adopt_bo_policy_template.docx", target, prepared.context)
            document = Document(target)
            table = next(table for table in document.tables if "PERSON 1" in table.rows[1].cells[0].text)
            self.assertEqual(len(table.rows), 4)
            self.assertEqual(table.rows[3].cells[0].text, "PERSON 3")
            self.assertEqual(table.rows[3].cells[2].text, "")

    def test_engagement_acknowledgement_uses_director_one(self):
        provider = FakeProvider([sample_row(directors=2, members=1)])
        with tempfile.TemporaryDirectory() as directory:
            result = generate_new_incorp_documents("EXAMPLE", Path(directory), provider=provider)
            self.assertEqual(result.status, "generated", result.issues)
            document = Document(next(path for path in result.output_paths if path.name == "Engagement Letter.docx"))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("PERSON 1", text)
            self.assertNotIn("PERSON 2", text)
            self.assertIn("20 Business Road", text)


if __name__ == "__main__":
    unittest.main()
